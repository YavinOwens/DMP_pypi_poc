#!/usr/bin/env python3
"""
RAG System for Enhancing Benchmark Reports
Uses Snowflake Vector Store with LangChain to enhance Word document reports
using knowledge from chunked PDF documents.

Architecture:
1. Extract text from Word documents (benchmark reports)
2. Chunk PDF documents for knowledge base
3. Store embeddings in Snowflake vector store
4. Use RAG to enhance reports with PDF knowledge
5. Generate enhanced reports with citations
"""

import sys
import json
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

# Configure logging - use smart configuration
try:
    from ..logging_config import configure_logging
    # Only configure if not already configured
    if not logging.getLogger().handlers:
        configure_logging(verbose=False)
except ImportError:
    # Fallback
    logging.basicConfig(
        level=logging.WARNING,  # Less verbose by default
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

logger = logging.getLogger(__name__)

# Try to import required libraries (optional - for PDF processing)
try:
    from langchain_community.document_loaders import PyPDFLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    logger.warning(f"LangChain not available for PDF processing: {e}")
    logger.info("PDF processing will use alternative methods")
    LANGCHAIN_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

try:
    from snowflake.snowpark import Session
    from snowflake.snowpark.functions import col, lit
    SNOWPARK_AVAILABLE = True
except ImportError:
    SNOWPARK_AVAILABLE = False
    logger.warning("Snowpark not available")

# Import project modules
try:
    from ..helpers import get_snowflake_session
except ImportError:
    logger.warning("Could not import get_snowflake_session from helpers")


class RAGSystem:
    """
    RAG System for enhancing benchmark reports using PDF knowledge base
    Processes documents from knowledgebase_docs/ folder as authoritative sources
    """
    
    def __init__(
        self,
        session: Optional[Any] = None,
        embedding_model: str = "snowflake-arctic-embed-l-v2.0",
        vector_store_table: str = "RAG_KNOWLEDGE_BASE",
        vector_store_schema: str = "PUBLIC",
        vector_store_database: Optional[str] = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        knowledge_base_dir: Optional[Path] = None
    ):
        """
        Initialize RAG System using Snowflake AI_EMBED function (recommended)
        
        Args:
            session: Snowflake session
            embedding_model: Snowflake embedding model name
                Options: 'e5-base-v2' (default, 768 dims),
                         'snowflake-arctic-embed-l-v2.0',
                         'snowflake-arctic-embed-l-v2.0-8k',
                         'nv-embed-qa-4', 'multilingual-e5-large'
            vector_store_table: Table name for vector store
            vector_store_schema: Schema for vector store
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.session = session or self._get_session()
        if not self.session:
            raise ValueError("Snowflake session required")
        
        self.embedding_model = embedding_model
        self.vector_store_table = vector_store_table
        self.vector_store_schema = vector_store_schema
        self.vector_store_database = vector_store_database  # Optional: override database
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Set knowledge base directory (default: knowledgebase_docs/ in package root or parent)
        if knowledge_base_dir:
            self.knowledge_base_dir = Path(knowledge_base_dir)
        else:
            # Try package root, then parent directory (for knowledgebase_docs/)
            package_root = Path(__file__).parent.parent.parent
            kb_dir = package_root / "knowledgebase_docs"
            if not kb_dir.exists():
                kb_dir = package_root.parent / "knowledgebase_docs"
            self.knowledge_base_dir = kb_dir
        
        # Ensure knowledge base directory exists
        self.knowledge_base_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Knowledge base directory: {self.knowledge_base_dir}")
        
        # Set up project-local temporary directory
        package_root = Path(__file__).parent.parent.parent
        self.tmp_dir = package_root / ".tmp"
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Temporary directory: {self.tmp_dir}")
        
        # Initialize components - ALWAYS use LangChain if available
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]  # Smart separators for better chunking
            )
            logger.info("✓ Using LangChain RecursiveCharacterTextSplitter for intelligent chunking")
        else:
            raise ImportError(
                "LangChain is required. Install with: pip install langchain langchain-community pypdf"
            )
        
        # Snowflake native embedding - no external model needed
        self._setup_vector_store()
    
    def _get_session(self):
        """Get Snowflake session"""
        try:
            return get_snowflake_session()
        except Exception as e:
            logger.error(f"Could not get Snowflake session: {e}")
            return None
    
    def _get_embedding_function(self) -> tuple:
        """
        Determine which Snowflake embedding function and model to use
        
        Returns:
            Tuple of (function_name, model_name, dimensions)
        """
        # Snowflake recommends using AI_EMBED (not EMBED_TEXT_768/1024)
        # AI_EMBED returns VECTOR type directly and handles dimensions automatically
        # Model dimensions:
        # - 'e5-base-v2': 768 dimensions
        # - 'snowflake-arctic-embed-l-v2.0': varies by model
        # - 'snowflake-arctic-embed-m-v1.5': varies by model
        
        # Use AI_EMBED with e5-base-v2 (768 dimensions) for compatibility
        # AI_EMBED is the recommended modern function
        return ('AI_EMBED', 'e5-base-v2', 768)
    
    def _generate_embedding_snowflake(self, text: str) -> List[float]:
        """
        Generate embedding using Snowflake AI_EMBED function (recommended)
        
        Uses AI_EMBED which is the modern recommended function (not EMBED_TEXT_768/1024)
        
        Args:
            text: Text to embed (should be reasonable length, truncated if too long)
            
        Returns:
            Embedding vector as list of floats
        """
        try:
            # Truncate text if too long (Snowflake functions may have limits)
            # AI_EMBED typically handles up to 8192 characters depending on model
            max_length = 8000  # Safe limit
            if len(text) > max_length:
                text = text[:max_length]
                logger.debug(f"Text truncated to {max_length} characters for embedding")
            
            # Escape single quotes for SQL
            text_escaped = text.replace("'", "''")
            
            # Determine which embedding function to use
            embed_func, model_name, dims = self._get_embedding_function()
            
            # Use Snowflake AI_EMBED function (recommended modern approach)
            # Syntax: AI_EMBED('model-name', 'text')
            # Returns VECTOR type directly
            # Note: AI_EMBED is in the default namespace, not SNOWFLAKE.CORTEX
            query = f"""
            SELECT AI_EMBED('{model_name}', '{text_escaped}') AS embedding
            """
            
            result = self.session.sql(query).collect()
            if result and len(result) > 0:
                # Access embedding exactly like snowflake_helpers.py line 770
                # result[0] is the first row, result[0][0] is the first column (embedding)
                row = result[0]
                if isinstance(row, (list, tuple)):
                    embedding = row[0]
                elif hasattr(row, '__getitem__'):
                    # Try index access
                    try:
                        embedding = row[0]
                    except (KeyError, IndexError):
                        # Try column name access
                        embedding = row.get("EMBEDDING") or row.get("embedding") if isinstance(row, dict) else None
                        if embedding is None:
                            # Try attribute access
                            embedding = getattr(row, "EMBEDDING", None) or getattr(row, "embedding", None)
                else:
                    # Try attribute or dict access
                    if isinstance(row, dict):
                        embedding = row.get("EMBEDDING") or row.get("embedding")
                    else:
                        embedding = getattr(row, "EMBEDDING", None) or getattr(row, "embedding", None)
                
                if embedding is None:
                    logger.warning("Could not extract embedding from result")
                    return None
                
                # Convert to list - match snowflake_helpers.py pattern
                if isinstance(embedding, list):
                    embedding_list = embedding
                elif hasattr(embedding, 'tolist'):
                    embedding_list = embedding.tolist()
                else:
                    try:
                        embedding_list = list(embedding)
                    except Exception as e:
                        logger.error(f"Could not convert embedding to list: {e}, type: {type(embedding)}")
                        return None
                
                # Verify dimension matches expected (critical for vector storage)
                expected_dim = dims
                if len(embedding_list) != expected_dim:
                    logger.warning(f"Embedding dimension mismatch: got {len(embedding_list)}, expected {expected_dim}. Adjusting...")
                    # Adjust dimension to match table definition
                    if len(embedding_list) > expected_dim:
                        embedding_list = embedding_list[:expected_dim]
                    else:
                        # Pad with zeros if too short (shouldn't happen, but safety check)
                        embedding_list = list(embedding_list) + [0.0] * (expected_dim - len(embedding_list))
                
                return embedding_list
            
            logger.warning("No embedding returned from Snowflake")
            return None
            
        except Exception as e:
            logger.error(f"Error generating embedding with AI_EMBED: {e}")
            # Try alternative: use Snowpark AI functions if available
            try:
                from snowflake.snowpark.functions import ai_embed, lit, col
                _, model_name, dims = self._get_embedding_function()
                
                # Use Snowpark's ai_embed function (recommended Snowpark approach)
                df = self.session.create_dataframe([[text[:8000]]], schema=["text"])
                result_df = df.select(ai_embed(model_name, col("text")).alias("embedding"))
                result = result_df.collect()
                
                if result and len(result) > 0:
                    embedding = result[0]["EMBEDDING"]
                    if isinstance(embedding, list):
                        embedding_list = embedding
                    elif hasattr(embedding, 'tolist'):
                        embedding_list = embedding.tolist()
                    else:
                        embedding_list = list(embedding)
                    
                    # Verify dimension
                    if len(embedding_list) != dims:
                        logger.warning(f"Alternative method dimension mismatch: got {len(embedding_list)}, expected {dims}")
                        if len(embedding_list) > dims:
                            embedding_list = embedding_list[:dims]
                        else:
                            embedding_list = list(embedding_list) + [0.0] * (dims - len(embedding_list))
                    
                    return embedding_list
            except Exception as e2:
                logger.debug(f"Alternative embedding method (Snowpark AI) also failed: {e2}")
            
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def _setup_vector_store(self):
        """Setup Snowflake vector store"""
        try:
            logger.info(f"Setting up vector store: {self.vector_store_schema}.{self.vector_store_table}")
            
            # Get current database and schema context
            try:
                current_db = self.session.sql("SELECT CURRENT_DATABASE()").collect()[0][0]
                current_schema = self.session.sql("SELECT CURRENT_SCHEMA()").collect()[0][0]
                logger.info(f"Current context: {current_db}.{current_schema}")
                
                # Use specified database or current database
                target_db = self.vector_store_database or current_db
                
                # Check if current database is read-only (like SNOWFLAKE_SAMPLE_DATA)
                if "SAMPLE_DATA" in current_db.upper() or current_db.upper() == "SNOWFLAKE_SAMPLE_DATA":
                    logger.warning(f"Current database '{current_db}' appears to be read-only")
                    if not self.vector_store_database:
                        logger.info("Attempting to use or create a writable database...")
                        # Try to use a common writable database name
                        target_db = os.getenv("SNOWFLAKE_RAG_DATABASE", "RAG_KNOWLEDGE_BASE")
                        try:
                            # Try to create database if it doesn't exist
                            self.session.sql(f"CREATE DATABASE IF NOT EXISTS {target_db}").collect()
                            logger.info(f"✓ Using/created database: {target_db}")
                        except Exception as db_error:
                            logger.warning(f"Could not create database {target_db}: {db_error}")
                            logger.info("Will attempt to use current database anyway...")
                            target_db = current_db
                
                # Always use current schema instead of hardcoded PUBLIC
                # This avoids schema permission issues
                if self.vector_store_schema == "PUBLIC":
                    self.vector_store_schema = current_schema
                    logger.info(f"Using current schema instead of PUBLIC: {self.vector_store_schema}")
                
                # Try to create schema if it doesn't exist (in target database)
                try:
                    self.session.sql(f"CREATE SCHEMA IF NOT EXISTS {target_db}.{self.vector_store_schema}").collect()
                    logger.info(f"✓ Schema {target_db}.{self.vector_store_schema} is ready")
                except Exception as schema_error:
                    logger.warning(f"Could not create schema {target_db}.{self.vector_store_schema}: {schema_error}")
                    logger.info("Will attempt to use existing schema...")
                
                # Update database reference
                self.vector_store_database = target_db
                
            except Exception as e:
                logger.warning(f"Could not determine current schema: {e}")
                # Fallback: use current database if specified database not available
                if not self.vector_store_database:
                    try:
                        self.vector_store_database = self.session.sql("SELECT CURRENT_DATABASE()").collect()[0][0]
                    except:
                        pass
            
            # Create vector store table if it doesn't exist
            self._create_vector_store_table()
            
            # Create ML training data table for logging low-quality results
            self._create_ml_training_table()
            
            logger.info("✓ Vector store setup complete")
        except Exception as e:
            logger.error(f"Failed to setup vector store: {e}")
            raise
    
    def _create_vector_store_table(self):
        """Create vector store table in Snowflake"""
        # Determine vector dimension based on embedding function
        _, _, vector_dim = self._get_embedding_function()
        
        # Use full database.schema.table path if database is specified
        table_path = f"{self.vector_store_database}.{self.vector_store_schema}.{self.vector_store_table}" if self.vector_store_database else f"{self.vector_store_schema}.{self.vector_store_table}"
        
        create_table_sql = f"""
        CREATE TABLE IF NOT EXISTS {table_path} (
            id VARCHAR(255) PRIMARY KEY,
            content TEXT,
            metadata VARIANT,
            embedding VECTOR(FLOAT, {vector_dim}),
            source VARCHAR(500),
            chunk_index INT,
            created_at TIMESTAMP_NTZ DEFAULT CURRENT_TIMESTAMP()
        )
        """
        try:
            self.session.sql(create_table_sql).collect()
            logger.info(f"✓ Vector store table created/verified: {self.vector_store_table} (dimension: {vector_dim})")
        except Exception as e:
            logger.warning(f"Could not create vector store table: {e}")
            logger.info("You may need to create it manually with appropriate permissions")
    
    def _create_ml_training_table(self):
        """Create ML training data table for logging low-quality retrieval results"""
        # Use same database/schema as vector store
        table_path = f"{self.vector_store_database}.{self.vector_store_schema}.RAG_ML_TRAINING_DATA" if self.vector_store_database else f"{self.vector_store_schema}.RAG_ML_TRAINING_DATA"
        
        create_table_sql = f"""
        CREATE TABLE IF NOT EXISTS {table_path} (
            training_id VARCHAR(255) PRIMARY KEY,
            session_id VARCHAR(255),
            section_name VARCHAR(500),
            query_text TEXT,
            query_embedding VECTOR(FLOAT, 768),
            retrieval_attempt INT,
            top_k INT,
            chunks_retrieved INT,
            avg_similarity FLOAT,
            min_similarity FLOAT,
            max_similarity FLOAT,
            quality_flag VARCHAR(50),
            quality_reason TEXT,
            retrieved_chunks VARIANT,
            validation_results VARIANT,
            enhancement_successful BOOLEAN,
            citation_validation_passed BOOLEAN,
            benchmark_context VARIANT,
            created_at TIMESTAMP_NTZ DEFAULT CURRENT_TIMESTAMP(),
            metadata VARIANT
        )
        """
        try:
            self.session.sql(create_table_sql).collect()
            logger.info(f"✓ ML training data table created/verified: RAG_ML_TRAINING_DATA")
        except Exception as e:
            logger.warning(f"Could not create ML training data table: {e}")
            logger.info("You may need to create it manually with appropriate permissions")
    
    def extract_text_from_word(self, docx_path: Path) -> str:
        """
        Extract text from Word document
        
        Args:
            docx_path: Path to .docx file
            
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for Word document extraction")
        
        logger.info(f"Extracting text from: {docx_path}")
        doc = Document(docx_path)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"✓ Extracted {len(full_text)} characters from Word document")
        return full_text
    
    def chunk_pdf(self, pdf_path: Path) -> List[Dict[str, Any]]:
        """
        Load and chunk PDF document using LangChain's intelligent text splitting
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of document chunks with metadata (as dicts)
        """
        if not LANGCHAIN_AVAILABLE:
            raise ImportError(
                "LangChain is required for PDF processing. "
                "Install with: pip install langchain langchain-community pypdf"
            )
        
        logger.info(f"Loading and chunking PDF using LangChain: {pdf_path}")
        logger.info(f"  PDF file size: {pdf_path.stat().st_size / (1024*1024):.2f} MB")
        
        # Use LangChain for PDF loading
        logger.info("  Loading PDF pages (this may take a while for large PDFs)...")
        loader = PyPDFLoader(str(pdf_path))
        documents = loader.load()
        
        logger.info(f"  ✓ Loaded {len(documents)} pages from PDF")
        
        # Add source metadata with authoritative flag
        for doc in documents:
            doc.metadata['source'] = str(pdf_path.name)
            doc.metadata['source_type'] = 'pdf'
            doc.metadata['authoritative'] = True
            doc.metadata['knowledge_base'] = True
        
        # Ensure text splitter is initialized (should be from __init__)
        if not self.text_splitter:
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]  # Smart separators for better chunking
            )
        
        # Split into chunks using LangChain's intelligent splitting
        chunks = self.text_splitter.split_documents(documents)
        
        logger.info(f"  Split into {len(chunks)} chunks using RecursiveCharacterTextSplitter")
        
        # Convert to dict format
        chunk_dicts = []
        for idx, chunk in enumerate(chunks):
            # Ensure chunk content is not too long for embedding
            content = chunk.page_content
            if len(content) > 8000:
                logger.warning(f"Chunk {idx} is {len(content)} chars, will be truncated during embedding")
            
            chunk_dicts.append({
                'page_content': content,
                'metadata': {
                    **chunk.metadata,
                    'chunk_index': idx
                }
            })
        
        logger.info(f"✓ Created {len(chunk_dicts)} chunks from PDF using LangChain")
        return chunk_dicts
    
    def store_pdf_in_vector_store(self, pdf_path: Path, max_chunks: Optional[int] = None, max_workers: Optional[int] = None) -> int:
        """
        Process PDF and store in Snowflake vector store
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Number of chunks stored
        """
        # Chunk PDF (works with or without LangChain)
        chunks = self.chunk_pdf(pdf_path)
        
        # Limit chunks if specified (for testing)
        if max_chunks and max_chunks > 0:
            chunks = chunks[:max_chunks]
            logger.info(f"  Limited to first {max_chunks} chunks for processing")
        
        # Generate embeddings and store with parallel processing
        total_chunks = len(chunks)
        
        # Determine parallelization settings
        if max_workers is None:
            max_workers = min(10, total_chunks)  # Use up to 10 workers, but not more than chunks
        else:
            max_workers = min(max_workers, total_chunks)  # Don't exceed chunk count
        use_parallel = total_chunks > 5  # Only parallelize if we have more than 5 chunks
        
        if use_parallel:
            estimated_time = (total_chunks / max_workers) * 1.5  # Parallel reduces time
            logger.info(f"Generating embeddings using AI_EMBED and storing {total_chunks} chunks...")
            logger.info(f"  Using parallel processing with {max_workers} workers")
            logger.info(f"  Estimated time: {estimated_time/60:.1f} minutes ({estimated_time:.0f} seconds)")
        else:
            estimated_time = total_chunks * 1.5  # ~1.5 seconds per chunk
            logger.info(f"Generating embeddings using AI_EMBED and storing {total_chunks} chunks...")
            logger.info(f"  Estimated time: {estimated_time/60:.1f} minutes ({estimated_time:.0f} seconds)")
        
        logger.info(f"  Progress will be shown every 10 chunks...")
        
        stored_count = 0
        failed_count = 0
        
        def process_chunk(chunk_data):
            """Process a single chunk: generate embedding and store"""
            idx, chunk = chunk_data
            try:
                # Get content (handle both dict and LangChain document formats)
                content = chunk.get('page_content') if isinstance(chunk, dict) else chunk.page_content
                metadata = chunk.get('metadata') if isinstance(chunk, dict) else chunk.metadata
                
                # Generate embedding using Snowflake native function
                # Note: Each thread uses the shared session (Snowflake sessions are thread-safe)
                embedding = self._generate_embedding_snowflake(content)
                
                if embedding is None:
                    return (idx, False, "Failed to generate embedding")
                
                # Store in Snowflake
                chunk_id = f"{pdf_path.stem}_{idx}"
                # Escape single quotes in content for SQL
                content_escaped = content.replace("'", "''").replace("\\", "\\\\")
                metadata_json = json.dumps(metadata).replace("'", "''")
                
                # Get expected dimension from function
                _, _, expected_dim = self._get_embedding_function()
                
                # Ensure embedding has correct dimension
                if len(embedding) != expected_dim:
                    if len(embedding) > expected_dim:
                        embedding = embedding[:expected_dim]
                    else:
                        embedding = list(embedding) + [0.0] * (expected_dim - len(embedding))
                
                # Convert embedding to JSON array format for VECTOR type
                embedding_json = json.dumps([float(x) for x in embedding])
                
                # Use full database.schema.table path if database is specified
                table_path = f"{self.vector_store_database}.{self.vector_store_schema}.{self.vector_store_table}" if self.vector_store_database else f"{self.vector_store_schema}.{self.vector_store_table}"
                
                # Use SELECT statement instead of VALUES to allow VECTOR casting
                insert_sql = f"""
                INSERT INTO {table_path}
                (id, content, metadata, embedding, source, chunk_index)
                SELECT 
                    '{chunk_id}',
                    '{content_escaped}',
                    PARSE_JSON('{metadata_json}'),
                    PARSE_JSON('{embedding_json}')::VECTOR(FLOAT, {expected_dim}),
                    '{pdf_path.name}',
                    {idx}
                """
                # Use shared session (Snowflake sessions are thread-safe for concurrent operations)
                self.session.sql(insert_sql).collect()
                return (idx, True, None)
            
            except Exception as e:
                return (idx, False, str(e))
        
        if use_parallel:
            # Parallel processing
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all tasks
                future_to_chunk = {
                    executor.submit(process_chunk, (idx, chunk)): (idx, chunk)
                    for idx, chunk in enumerate(chunks)
                }
                
                # Process completed tasks
                completed = 0
                for future in as_completed(future_to_chunk):
                    completed += 1
                    try:
                        idx, success, error = future.result()
                        if success:
                            stored_count += 1
                        else:
                            failed_count += 1
                            if error:
                                logger.warning(f"Failed to store chunk {idx}: {error}")
                        
                        # Progress logging every 10 chunks
                        if completed % 10 == 0:
                            logger.info(f"  Progress: {completed}/{total_chunks} chunks processed ({stored_count} stored, {failed_count} failed)...")
                    except Exception as e:
                        failed_count += 1
                        logger.warning(f"Error processing chunk: {e}")
        else:
            # Sequential processing (for small batches)
            for idx, chunk in enumerate(chunks):
                try:
                    # Progress logging every 10 chunks
                    if idx % 10 == 0 and idx > 0:
                        logger.info(f"  Progress: {idx}/{len(chunks)} chunks processed ({stored_count} stored)...")
                    
                    # Get content (handle both dict and LangChain document formats)
                    content = chunk.get('page_content') if isinstance(chunk, dict) else chunk.page_content
                    metadata = chunk.get('metadata') if isinstance(chunk, dict) else chunk.metadata
                    
                    # Generate embedding using Snowflake native function
                    if idx == 0:
                        logger.info(f"  Generating first embedding (chunk {idx+1}/{len(chunks)})...")
                    embedding = self._generate_embedding_snowflake(content)
                    
                    if embedding is None:
                        logger.warning(f"Failed to generate embedding for chunk {idx}, skipping")
                        failed_count += 1
                        continue
                    
                    # Store in Snowflake
                    chunk_id = f"{pdf_path.stem}_{idx}"
                    # Escape single quotes in content for SQL
                    content_escaped = content.replace("'", "''").replace("\\", "\\\\")
                    metadata_json = json.dumps(metadata).replace("'", "''")
                    
                    # Get expected dimension from function
                    _, _, expected_dim = self._get_embedding_function()
                    
                    # Ensure embedding has correct dimension
                    if len(embedding) != expected_dim:
                        logger.warning(f"Embedding dimension mismatch: got {len(embedding)}, expected {expected_dim}. Adjusting...")
                        if len(embedding) > expected_dim:
                            embedding = embedding[:expected_dim]
                        else:
                            embedding = list(embedding) + [0.0] * (expected_dim - len(embedding))
                    
                    # Convert embedding to JSON array format for VECTOR type
                    embedding_json = json.dumps([float(x) for x in embedding])
                    
                    # Use full database.schema.table path if database is specified
                    table_path = f"{self.vector_store_database}.{self.vector_store_schema}.{self.vector_store_table}" if self.vector_store_database else f"{self.vector_store_schema}.{self.vector_store_table}"
                    
                    # Use SELECT statement instead of VALUES to allow VECTOR casting
                    insert_sql = f"""
                    INSERT INTO {table_path}
                    (id, content, metadata, embedding, source, chunk_index)
                    SELECT 
                        '{chunk_id}',
                        '{content_escaped}',
                        PARSE_JSON('{metadata_json}'),
                        PARSE_JSON('{embedding_json}')::VECTOR(FLOAT, {expected_dim}),
                        '{pdf_path.name}',
                        {idx}
                    """
                    self.session.sql(insert_sql).collect()
                    stored_count += 1
                    
                    if (idx + 1) % 10 == 0:
                        logger.info(f"  ✓ Stored {idx + 1}/{len(chunks)} chunks ({stored_count} successful)")
                
                except Exception as e:
                    failed_count += 1
                    logger.warning(f"Failed to store chunk {idx}: {e}")
                    import traceback
                    logger.debug(traceback.format_exc())
        
        logger.info(f"✓ Stored {stored_count}/{len(chunks)} chunks in vector store ({failed_count} failed)")
        return stored_count
    
    def retrieve_relevant_chunks(
        self,
        query: str,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        Retrieve relevant chunks from vector store using semantic search
        
        Args:
            query: Search query
            top_k: Number of chunks to retrieve
            
        Returns:
            List of relevant chunks with metadata
        """
        logger.info(f"Retrieving top {top_k} relevant chunks for query...")
        
        # Generate query embedding using Snowflake native function
        query_embedding = self._generate_embedding_snowflake(query)
        
        if query_embedding is None:
            logger.error("Failed to generate query embedding")
            return []
        
        # Search in Snowflake using vector similarity
        # Using Snowflake's VECTOR_DOT_PRODUCT for cosine similarity
        # Ensure query embedding has correct dimension
        _, _, expected_dim = self._get_embedding_function()
        if len(query_embedding) != expected_dim:
            logger.warning(f"Query embedding dimension mismatch: got {len(query_embedding)}, expected {expected_dim}")
            if len(query_embedding) > expected_dim:
                query_embedding = query_embedding[:expected_dim]
            else:
                query_embedding = list(query_embedding) + [0.0] * (expected_dim - len(query_embedding))
        
        # Convert query embedding to SQL array format and cast to VECTOR
        embedding_values = [str(float(x)) for x in query_embedding]
        embedding_str = ','.join(embedding_values)
        
        # Use full database.schema.table path if database is specified
        table_path = f"{self.vector_store_database}.{self.vector_store_schema}.{self.vector_store_table}" if self.vector_store_database else f"{self.vector_store_schema}.{self.vector_store_table}"
        
        # Use VECTOR_COSINE_DISTANCE for similarity (lower = more similar)
        # Convert embedding to JSON array for comparison
        embedding_json = json.dumps([float(x) for x in query_embedding])
        
        search_sql = f"""
        SELECT 
            id,
            content,
            metadata,
            source,
            chunk_index,
            VECTOR_COSINE_DISTANCE(embedding, PARSE_JSON('{embedding_json}')::VECTOR(FLOAT, {expected_dim})) as distance
        FROM {table_path}
        ORDER BY distance ASC
        LIMIT {top_k}
        """
        
        try:
            results = self.session.sql(search_sql).collect()
            
            chunks = []
            for row in results:
                chunks.append({
                    'content': row['CONTENT'],
                    'metadata': json.loads(row['METADATA']) if isinstance(row['METADATA'], str) else row['METADATA'],
                    'source': row['SOURCE'],
                    'chunk_index': row['CHUNK_INDEX'],
                    'similarity': row['SIMILARITY']
                })
            
            logger.info(f"✓ Retrieved {len(chunks)} relevant chunks")
            return chunks
        
        except Exception as e:
            logger.error(f"Error retrieving chunks: {e}")
            # Fallback: try VECTOR_L2_DISTANCE or manual calculation
            logger.info("Trying alternative vector search syntax...")
            try:
                # Try VECTOR_L2_DISTANCE
                search_sql = f"""
                SELECT 
                    id,
                    content,
                    metadata,
                    source,
                    chunk_index,
                    VECTOR_L2_DISTANCE(embedding, PARSE_JSON('{embedding_json}')::VECTOR(FLOAT, {expected_dim})) as distance
                FROM {table_path}
                ORDER BY distance ASC
                LIMIT {top_k}
                """
                results = self.session.sql(search_sql).collect()
                
                chunks = []
                for row in results:
                    # Handle both dict-like and tuple-like row access
                    if isinstance(row, dict):
                        distance = row.get('DISTANCE', row.get('distance', 1.0))
                        content = row.get('CONTENT', row.get('content', ''))
                        metadata = row.get('METADATA', row.get('metadata', {}))
                        source = row.get('SOURCE', row.get('source', ''))
                        chunk_index = row.get('CHUNK_INDEX', row.get('chunk_index', 0))
                    else:
                        # Tuple-like access (index-based)
                        distance = row[5] if len(row) > 5 else 1.0
                        content = row[1] if len(row) > 1 else ''
                        metadata = row[2] if len(row) > 2 else {}
                        source = row[3] if len(row) > 3 else ''
                        chunk_index = row[4] if len(row) > 4 else 0
                    
                    similarity = 1.0 / (1.0 + float(distance))  # Convert L2 distance to similarity
                    
                    # Parse metadata if it's a string
                    if isinstance(metadata, str):
                        try:
                            metadata = json.loads(metadata)
                        except:
                            metadata = {}
                    
                    chunks.append({
                        'content': content,
                        'metadata': metadata,
                        'source': source,
                        'chunk_index': chunk_index,
                        'similarity': similarity,
                        'distance': distance
                    })
                
                logger.info(f"✓ Retrieved {len(chunks)} relevant chunks (using L2 distance)")
                return chunks
            except Exception as e2:
                logger.error(f"Alternative vector search also failed: {e2}")
                return []
    
    def enhance_report_section(
        self,
        section_text: str,
        section_name: str,
        use_rag: bool = True,
        top_k: int = 3,
        benchmark_context: Optional[Dict] = None,
        validate_citations: bool = True
    ) -> Dict[str, Any]:
        """
        Enhance a report section using RAG with validation checks
        
        Args:
            section_text: Original section text
            section_name: Name of the section
            use_rag: Whether to use RAG enhancement
            top_k: Number of relevant chunks to use
            benchmark_context: Context from benchmark results
            validate_citations: Whether to validate citations against PDF source
            
        Returns:
            Enhanced section with citations and validation results
        """
        if not use_rag:
            return {
                'original': section_text,
                'enhanced': section_text,
                'citations': [],
                'rag_used': False,
                'validation_passed': True
            }
        
        # Check if RAG system is properly initialized
        if not self._validate_rag_system():
            logger.warning("RAG system validation failed, returning original text")
            return {
                'original': section_text,
                'enhanced': section_text,
                'citations': [],
                'rag_used': False,
                'validation_passed': False,
                'validation_error': 'RAG system not properly initialized'
            }
        
        logger.info(f"Enhancing section: {section_name}")
        
        # Build enhanced query with benchmark context for better retrieval
        query_parts = [section_name]
        if benchmark_context:
            # Add relevant benchmark context to improve retrieval
            if 'best_model' in benchmark_context:
                query_parts.append(f"best model {benchmark_context['best_model']}")
            if 'top_3_models' in benchmark_context and benchmark_context['top_3_models']:
                top_models = ", ".join([m['model'] for m in benchmark_context['top_3_models'][:2]])
                query_parts.append(f"models {top_models}")
        
        # Combine section name, context, and preview for better retrieval
        enhanced_query = f"{' '.join(query_parts)}: {section_text[:300]}"
        
        # Retrieve relevant chunks with continuous validation
        retrieval_attempts = 0
        max_retrieval_attempts = 3
        relevant_chunks = []
        quality_metrics = []
        
        while retrieval_attempts < max_retrieval_attempts and not relevant_chunks:
            retrieval_attempts += 1
            logger.info(f"Retrieval attempt {retrieval_attempts}/{max_retrieval_attempts} for section: {section_name}")
            
            relevant_chunks = self.retrieve_relevant_chunks(
                query=enhanced_query,
                top_k=top_k
            )
            
            # Validate retrieval quality
            if relevant_chunks:
                similarities = [chunk.get('similarity', 0) for chunk in relevant_chunks]
                avg_similarity = sum(similarities) / len(similarities)
                min_similarity = min(similarities)
                max_similarity = max(similarities)
                
                # Store quality metrics for logging
                quality_metrics.append({
                    'attempt': retrieval_attempts,
                    'chunks_retrieved': len(relevant_chunks),
                    'avg_similarity': avg_similarity,
                    'min_similarity': min_similarity,
                    'max_similarity': max_similarity
                })
                
                if avg_similarity < 0.3:  # Low similarity threshold
                    logger.warning(f"Low average similarity ({avg_similarity:.3f}), retrying retrieval...")
                    # Log low-quality retrieval for ML training
                    self._log_retrieval_for_training(
                        section_name=section_name,
                        query_text=enhanced_query,
                        retrieval_attempt=retrieval_attempts,
                        top_k=top_k,
                        chunks_retrieved=relevant_chunks,
                        quality_flag='LOW_QUALITY',
                        quality_reason=f'Average similarity {avg_similarity:.3f} below threshold 0.3',
                        benchmark_context=benchmark_context
                    )
                    relevant_chunks = []  # Retry
                    continue
                else:
                    logger.info(f"✓ Retrieval validated: {len(relevant_chunks)} chunks with avg similarity {avg_similarity:.3f}")
                    break
        
        if not relevant_chunks:
            logger.warning("No relevant chunks found after validation attempts, returning original text")
            # Log failed retrieval for ML training
            self._log_retrieval_for_training(
                section_name=section_name,
                query_text=enhanced_query,
                retrieval_attempt=retrieval_attempts,
                top_k=top_k,
                chunks_retrieved=[],
                quality_flag='FAILED',
                quality_reason='No relevant chunks found after all retrieval attempts',
                benchmark_context=benchmark_context
            )
            return {
                'original': section_text,
                'enhanced': section_text,
                'citations': [],
                'rag_used': False,
                'validation_passed': False,
                'validation_error': 'No relevant chunks found after validation'
            }
        
        # Build context from relevant chunks with authoritative source formatting
        context_parts = []
        for chunk in relevant_chunks:
            source_name = chunk['source'].replace('.pdf', '').replace('_', ' ')
            context_parts.append(
                f"[Authoritative Source: {source_name}, Chunk {chunk['chunk_index']}]\n{chunk['content']}"
            )
        context = "\n\n".join(context_parts)
        
        # Build benchmark context summary for prompt
        benchmark_summary = ""
        if benchmark_context:
            if 'best_model' in benchmark_context:
                benchmark_summary += f"\nBest Model: {benchmark_context['best_model']}"
            if 'top_3_models' in benchmark_context and benchmark_context['top_3_models']:
                benchmark_summary += "\nTop Models: " + ", ".join([m['model'] for m in benchmark_context['top_3_models'][:3]])
        
        # Create enhancement prompt with authoritative source emphasis and benchmark context
        enhancement_prompt = f"""You are a technical writing assistant helping to enhance a benchmark report section using authoritative knowledge base sources.

Original Section ({section_name}):
{section_text}
{benchmark_summary}

Relevant Knowledge Base Context (Authoritative Sources):
{context}

Task:
1. Review the original section carefully, including the benchmark context
2. Enhance it by incorporating relevant information from the authoritative knowledge base context
3. Reference these knowledge base documents as authoritative sources (e.g., DAMA DMBOK 2nd Edition)
4. Paraphrase and integrate knowledge naturally while maintaining technical accuracy
5. Add citations using PROTECTED citation format: {{CITE: Document Name}} for inline references
6. Use full document name in citations (e.g., {{CITE: DAMA DMBOK 2nd Edition}} not just "DMBOK")
7. Maintain the professional tone suitable for technology and data management professionals
8. Do NOT change the core findings, metrics, or data - only enhance explanations, context, and add authoritative references
9. Use phrases like "According to industry standards {{CITE: DAMA DMBOK 2nd Edition}}" when referencing authoritative sources
10. Ensure all enhancements align with industry best practices from the knowledge base
11. Connect knowledge base insights to the specific benchmark findings and model performance data
12. CRITICAL: Use {{CITE: Document Name}} format for ALL citations - this format is protected from markdown cleaning
13. GRAMMAR: Use proper English grammar, spelling, and punctuation throughout. Ensure all sentences are grammatically correct with proper subject-verb agreement, consistent verb tenses, and complete sentence structure

Enhanced Section:"""
        
        # Use Snowflake Cortex AI to generate enhanced section
        try:
            enhanced_text = self._generate_with_cortex(enhancement_prompt)
            
            # Extract citations with authoritative source information
            citations = []
            validation_results = []
            
            for chunk in relevant_chunks:
                source_name = chunk['source']
                # Clean up source name (remove .pdf extension, format nicely)
                source_display = source_name.replace('.pdf', '').replace('_', ' ')
                
                citation_info = {
                    'source': source_name,
                    'source_display': source_display,
                    'chunk_index': chunk['chunk_index'],
                    'relevance': chunk.get('similarity', 0),
                    'authoritative': True,  # All knowledge base sources are authoritative
                    'knowledge_base': True,
                    'chunk_content': chunk.get('content', '')[:200]  # Store snippet for validation
                }
                
                # Validate citation against PDF source if requested
                if validate_citations:
                    validation_result = self._validate_citation_against_source(
                        citation_info, 
                        enhanced_text
                    )
                    citation_info['validation'] = validation_result
                    validation_results.append(validation_result)
                
                citations.append(citation_info)
            
            # Overall validation status
            validation_passed = all(
                result.get('is_valid', False) for result in validation_results
            ) if validation_results else True
            
            if not validation_passed:
                logger.warning(f"Citation validation found issues for section: {section_name}")
                for i, result in enumerate(validation_results):
                    if not result.get('is_valid', False):
                        logger.warning(f"  Citation {i+1}: {result.get('issue', 'Unknown issue')}")
            
            # Log successful retrieval for ML training (even if validation had issues)
            if relevant_chunks:
                similarities = [chunk.get('similarity', 0) for chunk in relevant_chunks]
                avg_similarity = sum(similarities) / len(similarities) if similarities else 0
                min_similarity = min(similarities) if similarities else 0
                max_similarity = max(similarities) if similarities else 0
                
                # Determine quality flag
                if avg_similarity < 0.5:
                    quality_flag = 'MEDIUM_QUALITY'
                    quality_reason = f'Average similarity {avg_similarity:.3f} below optimal threshold 0.5'
                else:
                    quality_flag = 'HIGH_QUALITY'
                    quality_reason = f'Average similarity {avg_similarity:.3f} meets quality standards'
                
                self._log_retrieval_for_training(
                    section_name=section_name,
                    query_text=enhanced_query,
                    retrieval_attempt=1,  # Final successful attempt
                    top_k=top_k,
                    chunks_retrieved=relevant_chunks,
                    quality_flag=quality_flag,
                    quality_reason=quality_reason,
                    validation_results=validation_results,
                    enhancement_successful=True,
                    citation_validation_passed=validation_passed,
                    benchmark_context=benchmark_context
                )
            
            return {
                'original': section_text,
                'enhanced': enhanced_text,
                'citations': citations,
                'rag_used': True,
                'context_chunks': len(relevant_chunks),
                'validation_passed': validation_passed,
                'validation_results': validation_results
            }
        
        except Exception as e:
            logger.error(f"Error enhancing section: {e}")
            return {
                'original': section_text,
                'enhanced': section_text,
                'citations': [],
                'rag_used': False,
                'error': str(e)
            }
    
    def _validate_rag_system(self) -> bool:
        """Validate that RAG system is properly initialized and accessible"""
        try:
            # Check if session is available
            if not self.session:
                logger.error("RAG validation failed: No Snowflake session")
                return False
            
            # Check if vector store table exists and is accessible
            table_path = f"{self.vector_store_database}.{self.vector_store_schema}.{self.vector_store_table}" if self.vector_store_database else f"{self.vector_store_schema}.{self.vector_store_table}"
            
            check_query = f"SELECT COUNT(*) as count FROM {table_path} LIMIT 1"
            result = self.session.sql(check_query).collect()
            
            if result:
                count = result[0][0] if isinstance(result[0], (list, tuple)) else result[0].get('COUNT', 0)
                if count == 0:
                    logger.warning(f"RAG validation: Vector store table exists but is empty ({table_path})")
                    return False
                logger.info(f"✓ RAG system validated: {count} chunks available in vector store")
                return True
            else:
                logger.error("RAG validation failed: Could not query vector store")
                return False
                
        except Exception as e:
            logger.error(f"RAG validation error: {e}")
            return False
    
    def _validate_citation_against_source(
        self, 
        citation_info: Dict[str, Any], 
        enhanced_text: str
    ) -> Dict[str, Any]:
        """
        Validate that citations in enhanced text actually reference content from the source PDF
        
        Args:
            citation_info: Citation metadata including source, chunk_index, chunk_content
            enhanced_text: The enhanced text that should contain citations
            
        Returns:
            Validation result with is_valid flag and any issues found
        """
        try:
            source_name = citation_info.get('source', '')
            chunk_content = citation_info.get('chunk_content', '')
            chunk_index = citation_info.get('chunk_index', 0)
            
            # Check if citation format appears in enhanced text
            citation_patterns = [
                f"{{CITE: {source_name.replace('.pdf', '')}}}",
                f"{{{{CITE: {source_name.replace('.pdf', '')}}}}}",
                source_name.replace('.pdf', ''),
                "DAMA DMBOK"
            ]
            
            citation_found = any(pattern in enhanced_text for pattern in citation_patterns)
            
            # Check if key terms from chunk appear in enhanced text (basic validation)
            if chunk_content:
                # Extract key terms from chunk (first 50 words)
                chunk_words = set(chunk_content.lower().split()[:50])
                enhanced_words = set(enhanced_text.lower().split())
                
                # Check for overlap of meaningful terms (longer than 4 chars)
                meaningful_chunk_terms = {w for w in chunk_words if len(w) > 4}
                meaningful_enhanced_terms = {w for w in enhanced_words if len(w) > 4}
                
                overlap = meaningful_chunk_terms.intersection(meaningful_enhanced_terms)
                overlap_ratio = len(overlap) / len(meaningful_chunk_terms) if meaningful_chunk_terms else 0
                
                # Validation passes if citation found and reasonable term overlap
                is_valid = citation_found and overlap_ratio > 0.1
                
                if not is_valid:
                    issue = f"Citation validation issue: "
                    if not citation_found:
                        issue += "Citation format not found in enhanced text. "
                    if overlap_ratio <= 0.1:
                        issue += f"Low term overlap ({overlap_ratio:.2%}) between source chunk and enhanced text."
                else:
                    issue = None
                
                return {
                    'is_valid': is_valid,
                    'citation_found': citation_found,
                    'term_overlap_ratio': overlap_ratio,
                    'source': source_name,
                    'chunk_index': chunk_index,
                    'issue': issue
                }
            else:
                return {
                    'is_valid': citation_found,
                    'citation_found': citation_found,
                    'term_overlap_ratio': 0.0,
                    'source': source_name,
                    'chunk_index': chunk_index,
                    'issue': 'No chunk content available for validation' if not citation_found else None
                }
                
        except Exception as e:
            logger.error(f"Citation validation error: {e}")
            return {
                'is_valid': False,
                'citation_found': False,
                'term_overlap_ratio': 0.0,
                'source': citation_info.get('source', 'unknown'),
                'chunk_index': citation_info.get('chunk_index', 0),
                'issue': f'Validation error: {str(e)}'
            }
    
    def _log_retrieval_for_training(
        self,
        section_name: str,
        query_text: str,
        retrieval_attempt: int,
        top_k: int,
        chunks_retrieved: List[Dict],
        quality_flag: str,
        quality_reason: str,
        benchmark_context: Optional[Dict] = None,
        validation_results: Optional[List[Dict]] = None,
        enhancement_successful: bool = False,
        citation_validation_passed: bool = True
    ):
        """
        Log retrieval results to ML training data table for continuous learning
        
        Args:
            section_name: Name of the section being enhanced
            query_text: The query text used for retrieval
            retrieval_attempt: Attempt number (1, 2, 3, etc.)
            top_k: Number of chunks requested
            chunks_retrieved: List of retrieved chunks with metadata
            quality_flag: Quality assessment (HIGH_QUALITY, MEDIUM_QUALITY, LOW_QUALITY, FAILED)
            quality_reason: Explanation of quality assessment
            benchmark_context: Context from benchmark results
            validation_results: Citation validation results
            enhancement_successful: Whether enhancement was successful
            citation_validation_passed: Whether citation validation passed
        """
        try:
            import uuid
            from datetime import datetime
            
            # Generate unique training ID
            training_id = str(uuid.uuid4())
            
            # Generate session ID (use timestamp-based for grouping)
            session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Calculate similarity metrics
            similarities = [chunk.get('similarity', 0) for chunk in chunks_retrieved] if chunks_retrieved else []
            avg_similarity = sum(similarities) / len(similarities) if similarities else 0.0
            min_similarity = min(similarities) if similarities else 0.0
            max_similarity = max(similarities) if similarities else 0.0
            
            # Generate query embedding for storage
            query_embedding = self._generate_embedding_snowflake(query_text[:8000])  # Limit text length
            if query_embedding:
                # Convert to JSON array for VECTOR type
                embedding_json = json.dumps([float(x) for x in query_embedding])
            else:
                embedding_json = None
            
            # Prepare chunks data (limit size for storage)
            chunks_data = []
            for chunk in chunks_retrieved[:top_k]:  # Limit to top_k
                chunks_data.append({
                    'source': chunk.get('source', ''),
                    'chunk_index': chunk.get('chunk_index', 0),
                    'similarity': chunk.get('similarity', 0),
                    'distance': chunk.get('distance', 0),
                    'content_preview': chunk.get('content', '')[:500]  # Limit content size
                })
            
            # Prepare table path
            table_path = f"{self.vector_store_database}.{self.vector_store_schema}.RAG_ML_TRAINING_DATA" if self.vector_store_database else f"{self.vector_store_schema}.RAG_ML_TRAINING_DATA"
            
            # Prepare JSON strings for SQL (escape single quotes)
            validation_results_json = json.dumps(validation_results).replace("'", "''") if validation_results else None
            benchmark_context_json = json.dumps(benchmark_context).replace("'", "''") if benchmark_context else None
            chunks_metadata_json = json.dumps({"chunks": chunks_data}).replace("'", "''")
            
            # Build INSERT statement
            insert_sql = f"""
            INSERT INTO {table_path} (
                training_id, session_id, section_name, query_text, query_embedding,
                retrieval_attempt, top_k, chunks_retrieved, avg_similarity,
                min_similarity, max_similarity, quality_flag, quality_reason,
                validation_results, enhancement_successful, citation_validation_passed,
                benchmark_context, metadata
            ) VALUES (
                '{training_id}',
                '{session_id}',
                '{section_name.replace("'", "''")}',
                '{query_text[:10000].replace("'", "''")}',
                {f"PARSE_JSON('{embedding_json}')::VECTOR(FLOAT, 768)" if embedding_json else "NULL"},
                {retrieval_attempt},
                {top_k},
                {len(chunks_retrieved)},
                {avg_similarity},
                {min_similarity},
                {max_similarity},
                '{quality_flag}',
                '{quality_reason.replace("'", "''")}',
                {f"PARSE_JSON('{validation_results_json}')" if validation_results_json else "NULL"},
                {str(enhancement_successful).upper()},
                {str(citation_validation_passed).upper()},
                {f"PARSE_JSON('{benchmark_context_json}')" if benchmark_context_json else "NULL"},
                PARSE_JSON('{chunks_metadata_json}')
            )
            """
            
            self.session.sql(insert_sql).collect()
            logger.info(f"✓ Logged retrieval data for ML training (ID: {training_id[:8]}..., Quality: {quality_flag})")
            
        except Exception as e:
            # Don't fail the main workflow if logging fails
            logger.warning(f"Could not log retrieval data for ML training: {e}")
    
    def _generate_with_cortex(self, prompt: str, model: str = "claude-3-5-sonnet") -> str:
        """Generate text using Snowflake Cortex AI"""
        try:
            from ..benchmark import execute_cortex_query
            return execute_cortex_query(self.session, model, prompt)
        except Exception as e:
            logger.error(f"Error generating with Cortex: {e}")
            raise
    
    def get_training_data_stats(self, days: int = 30) -> Dict[str, Any]:
        """
        Get statistics about training data for analysis
        
        Args:
            days: Number of days to look back
            
        Returns:
            Dictionary with statistics
        """
        try:
            table_path = f"{self.vector_store_database}.{self.vector_store_schema}.RAG_ML_TRAINING_DATA" if self.vector_store_database else f"{self.vector_store_schema}.RAG_ML_TRAINING_DATA"
            
            stats_sql = f"""
            SELECT 
                COUNT(*) as total_records,
                COUNT(DISTINCT section_name) as unique_sections,
                AVG(avg_similarity) as avg_similarity_overall,
                AVG(CASE WHEN quality_flag = 'HIGH_QUALITY' THEN 1 ELSE 0 END) * 100 as high_quality_pct,
                AVG(CASE WHEN quality_flag = 'LOW_QUALITY' THEN 1 ELSE 0 END) * 100 as low_quality_pct,
                AVG(CASE WHEN quality_flag = 'FAILED' THEN 1 ELSE 0 END) * 100 as failed_pct,
                AVG(CASE WHEN enhancement_successful THEN 1 ELSE 0 END) * 100 as enhancement_success_pct,
                AVG(CASE WHEN citation_validation_passed THEN 1 ELSE 0 END) * 100 as citation_validation_pct
            FROM {table_path}
            WHERE created_at >= DATEADD(day, -{days}, CURRENT_TIMESTAMP())
            """
            
            result = self.session.sql(stats_sql).collect()
            if result:
                row = result[0]
                return {
                    'total_records': row[0] if isinstance(row, (list, tuple)) else row.get('TOTAL_RECORDS', 0),
                    'unique_sections': row[1] if isinstance(row, (list, tuple)) else row.get('UNIQUE_SECTIONS', 0),
                    'avg_similarity_overall': float(row[2]) if isinstance(row, (list, tuple)) else float(row.get('AVG_SIMILARITY_OVERALL', 0)),
                    'high_quality_pct': float(row[3]) if isinstance(row, (list, tuple)) else float(row.get('HIGH_QUALITY_PCT', 0)),
                    'low_quality_pct': float(row[4]) if isinstance(row, (list, tuple)) else float(row.get('LOW_QUALITY_PCT', 0)),
                    'failed_pct': float(row[5]) if isinstance(row, (list, tuple)) else float(row.get('FAILED_PCT', 0)),
                    'enhancement_success_pct': float(row[6]) if isinstance(row, (list, tuple)) else float(row.get('ENHANCEMENT_SUCCESS_PCT', 0)),
                    'citation_validation_pct': float(row[7]) if isinstance(row, (list, tuple)) else float(row.get('CITATION_VALIDATION_PCT', 0))
                }
            return {}
        except Exception as e:
            logger.error(f"Error getting training data stats: {e}")
            return {}
    
    def get_low_quality_queries(self, limit: int = 100, days: int = 30) -> List[Dict[str, Any]]:
        """
        Get low-quality retrieval queries for ML training analysis
        
        Args:
            limit: Maximum number of records to return
            days: Number of days to look back
            
        Returns:
            List of low-quality query records
        """
        try:
            table_path = f"{self.vector_store_database}.{self.vector_store_schema}.RAG_ML_TRAINING_DATA" if self.vector_store_database else f"{self.vector_store_schema}.RAG_ML_TRAINING_DATA"
            
            query_sql = f"""
            SELECT 
                training_id,
                section_name,
                query_text,
                avg_similarity,
                quality_flag,
                quality_reason,
                chunks_retrieved,
                created_at
            FROM {table_path}
            WHERE quality_flag IN ('LOW_QUALITY', 'FAILED', 'MEDIUM_QUALITY')
            AND created_at >= DATEADD(day, -{days}, CURRENT_TIMESTAMP())
            ORDER BY avg_similarity ASC, created_at DESC
            LIMIT {limit}
            """
            
            results = self.session.sql(query_sql).collect()
            records = []
            for row in results:
                if isinstance(row, (list, tuple)):
                    records.append({
                        'training_id': row[0],
                        'section_name': row[1],
                        'query_text': row[2],
                        'avg_similarity': float(row[3]) if row[3] else 0,
                        'quality_flag': row[4],
                        'quality_reason': row[5],
                        'chunks_retrieved': row[6],
                        'created_at': row[7]
                    })
                else:
                    records.append({
                        'training_id': row.get('TRAINING_ID', ''),
                        'section_name': row.get('SECTION_NAME', ''),
                        'query_text': row.get('QUERY_TEXT', ''),
                        'avg_similarity': float(row.get('AVG_SIMILARITY', 0)),
                        'quality_flag': row.get('QUALITY_FLAG', ''),
                        'quality_reason': row.get('QUALITY_REASON', ''),
                        'chunks_retrieved': row.get('CHUNKS_RETRIEVED', 0),
                        'created_at': row.get('CREATED_AT', '')
                    })
            return records
        except Exception as e:
            logger.error(f"Error getting low-quality queries: {e}")
            return []
    
    def enhance_full_report(
        self,
        docx_path: Path,
        output_path: Optional[Path] = None,
        enhance_sections: Optional[List[str]] = None
    ) -> Path:
        """
        Enhance full Word document report using RAG
        
        Args:
            docx_path: Path to original Word document
            output_path: Path for enhanced document (optional)
            enhance_sections: List of section names to enhance (None = all)
            
        Returns:
            Path to enhanced document
        """
        logger.info(f"Enhancing full report: {docx_path}")
        
        # Extract text from Word document
        full_text = self.extract_text_from_word(docx_path)
        
        # Split into sections (simple approach - can be enhanced)
        sections = self._split_into_sections(full_text)
        
        # Enhance each section
        enhanced_sections = {}
        for section_name, section_text in sections.items():
            if enhance_sections and section_name not in enhance_sections:
                enhanced_sections[section_name] = {
                    'original': section_text,
                    'enhanced': section_text,
                    'citations': [],
                    'rag_used': False
                }
            else:
                enhanced_sections[section_name] = self.enhance_report_section(
                    section_text, section_name
                )
        
        # Create enhanced Word document
        if not output_path:
            output_path = docx_path.parent / f"{docx_path.stem}_enhanced_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        self._create_enhanced_document(docx_path, enhanced_sections, output_path)
        
        logger.info(f"✓ Enhanced report created: {output_path}")
        return output_path
    
    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """Split text into sections based on headings"""
        sections = {}
        current_section = "Introduction"
        current_text = []
        
        lines = text.split('\n')
        for line in lines:
            # Detect headings (all caps, or lines with #, or short lines)
            if (line.isupper() and len(line.split()) < 10) or line.startswith('#'):
                if current_text:
                    sections[current_section] = '\n'.join(current_text)
                current_section = line.strip('#').strip()
                current_text = []
            else:
                current_text.append(line)
        
        if current_text:
            sections[current_section] = '\n'.join(current_text)
        
        return sections
    
    def _create_enhanced_document(
        self,
        original_path: Path,
        enhanced_sections: Dict[str, Dict[str, Any]],
        output_path: Path
    ):
        """Create enhanced Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required")
        
        doc = Document()
        
        # Title
        doc.add_heading('Enhanced Benchmark Report (RAG-Enhanced)', 0)
        doc.add_paragraph(f'Original Report: {original_path.name}')
        doc.add_paragraph(f'Enhanced: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        # Add sections
        for section_name, section_data in enhanced_sections.items():
            doc.add_heading(section_name, level=1)
            
            if section_data.get('rag_used'):
                doc.add_paragraph('(Enhanced with RAG knowledge base)', style='Intense Quote')
            
            doc.add_paragraph(section_data['enhanced'])
            
            # Add citations if available (with authoritative source formatting)
            if section_data.get('citations'):
                doc.add_paragraph('')
                doc.add_heading('Authoritative References', level=2)
                doc.add_paragraph(
                    'The following authoritative knowledge base sources were referenced in this section:',
                    style='Intense Quote'
                )
                for citation in section_data['citations']:
                    source_display = citation.get('source_display', citation['source'])
                    doc.add_paragraph(
                        f"• {source_display} (Chunk {citation['chunk_index']})",
                        style='List Bullet'
                    )
            
            doc.add_paragraph('')
        
        doc.save(str(output_path))


def process_knowledge_base_folder(rag: 'RAGSystem', max_chunks: Optional[int] = None, max_workers: Optional[int] = None) -> int:
    """
    Process all PDFs in the knowledgebase_docs/ folder
    
    Args:
        rag: RAGSystem instance
        
    Returns:
        Total number of chunks stored
    """
    total_chunks = 0
    
    if not rag.knowledge_base_dir.exists():
        logger.warning(f"Knowledge base directory not found: {rag.knowledge_base_dir}")
        return 0
    
    # Find all PDF files
    pdf_files = list(rag.knowledge_base_dir.glob("*.pdf"))
    
    if not pdf_files:
        logger.warning(f"No PDF files found in {rag.knowledge_base_dir}")
        return 0
    
    logger.info(f"Found {len(pdf_files)} PDF file(s) in knowledge base directory")
    
    for pdf_path in pdf_files:
        logger.info(f"\nProcessing: {pdf_path.name}")
        try:
            chunks_stored = rag.store_pdf_in_vector_store(pdf_path, max_chunks=max_chunks, max_workers=max_workers)
            total_chunks += chunks_stored
            logger.info(f"✓ Stored {chunks_stored} chunks from {pdf_path.name}")
        except Exception as e:
            logger.error(f"Failed to process {pdf_path.name}: {e}")
    
    return total_chunks


def main():
    """Example usage of RAG System"""
    import argparse
    import traceback
    
    parser = argparse.ArgumentParser(
        description='RAG System for enhancing benchmark reports using authoritative knowledge base',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process all PDFs in knowledgebase_docs/ folder
  python rag_system.py --process-knowledge-base
  
  # Process a specific PDF
  python rag_system.py --store-pdf --pdf "knowledgebase_docs/DAMA DMBOK 2nd Edition.pdf"
  
  # Enhance a benchmark report
  python rag_system.py --enhance --docx tests/20260107_141843/model_benchmark_report_20260107_144512.docx
  
  # Process knowledge base and enhance report
  python rag_system.py --process-knowledge-base --enhance --docx tests/20260107_141843/model_benchmark_report_20260107_144512.docx
        """
    )
    
    parser.add_argument('--pdf', type=str, help='Path to specific PDF knowledge base file')
    parser.add_argument('--docx', type=str, help='Path to Word document to enhance')
    parser.add_argument('--store-pdf', action='store_true', help='Store specific PDF in vector store')
    parser.add_argument('--process-knowledge-base', action='store_true', help='Process all PDFs in knowledgebase_docs/ folder')
    parser.add_argument('--enhance', action='store_true', help='Enhance Word document')
    parser.add_argument('--output', type=str, help='Output path for enhanced document')
    parser.add_argument('--knowledge-base-dir', type=str, help='Path to knowledge base directory (default: knowledgebase_docs/)')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose logging')
    parser.add_argument('--max-chunks', type=int, help='Maximum number of chunks to process (for testing, default: all)')
    parser.add_argument('--rag-database', type=str, help='Database name for vector store (default: auto-detect or create RAG_KNOWLEDGE_BASE)')
    parser.add_argument('--rag-schema', type=str, help='Schema name for vector store (default: current schema or PUBLIC)')
    parser.add_argument('--max-workers', type=int, help='Maximum number of parallel workers for indexing (default: 10)')
    
    args = parser.parse_args()
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.setLevel(logging.DEBUG)
    
    logger.info("=" * 80)
    logger.info("RAG SYSTEM - Starting")
    logger.info("=" * 80)
    
    # Initialize RAG system
    logger.info("Initializing RAG system...")
    try:
        knowledge_base_dir = Path(args.knowledge_base_dir) if args.knowledge_base_dir else None
        rag_database = args.rag_database or os.getenv("SNOWFLAKE_RAG_DATABASE")
        rag_schema = args.rag_schema or os.getenv("SNOWFLAKE_RAG_SCHEMA", "PUBLIC")
        
        logger.info(f"Knowledge base directory: {knowledge_base_dir}")
        if rag_database:
            logger.info(f"Vector store database: {rag_database}")
        if rag_schema:
            logger.info(f"Vector store schema: {rag_schema}")
        
        rag = RAGSystem(
            knowledge_base_dir=knowledge_base_dir,
            vector_store_database=rag_database,
            vector_store_schema=rag_schema
        )
        logger.info("✓ RAG system initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize RAG system: {e}")
        logger.error(traceback.format_exc())
        return 1
    
    # Process knowledge base folder if requested
    if args.process_knowledge_base:
        logger.info("=" * 80)
        logger.info("PROCESSING KNOWLEDGE BASE FOLDER")
        logger.info("=" * 80)
        try:
            total_chunks = process_knowledge_base_folder(rag, max_chunks=args.max_chunks, max_workers=getattr(args, 'max_workers', None))
            logger.info(f"\n✓ Total chunks stored from knowledge base: {total_chunks}")
            if total_chunks == 0:
                logger.warning("No chunks were stored. Check if PDFs exist in knowledge base directory.")
        except Exception as e:
            logger.error(f"Error processing knowledge base: {e}")
            logger.error(traceback.format_exc())
            return 1
    
    # Store specific PDF if requested
    elif args.store_pdf and args.pdf:
        pdf_path = Path(args.pdf)
        if not pdf_path.exists():
            logger.error(f"PDF not found: {pdf_path}")
            return 1
        
        logger.info("Storing PDF in vector store...")
        stored = rag.store_pdf_in_vector_store(pdf_path, max_chunks=args.max_chunks, max_workers=getattr(args, 'max_workers', None))
        logger.info(f"✓ Stored {stored} chunks from PDF")
    
    # Enhance document if requested
    if args.enhance and args.docx:
        docx_path = Path(args.docx)
        if not docx_path.exists():
            logger.error(f"Word document not found: {docx_path}")
            return 1
        
        logger.info("=" * 80)
        logger.info("ENHANCING BENCHMARK REPORT WITH RAG")
        logger.info("=" * 80)
        output_path = Path(args.output) if args.output else None
        enhanced_path = rag.enhance_full_report(docx_path, output_path)
        logger.info(f"\n✓ Enhanced report created: {enhanced_path}")
        logger.info("\nThe enhanced report includes:")
        logger.info("  - Original content preserved")
        logger.info("  - Authoritative references from knowledge base")
        logger.info("  - Citations to knowledge base documents")
        logger.info("  - Professional enhancements aligned with industry standards")
    
    if not any([args.process_knowledge_base, args.store_pdf, args.enhance]):
        logger.warning("No action specified. Use --process-knowledge-base, --store-pdf, or --enhance")
        parser.print_help()
        return 1
    
    return 0


if __name__ == "__main__":
    sys.exit(main())
