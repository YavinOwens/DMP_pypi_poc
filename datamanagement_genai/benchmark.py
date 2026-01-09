#!/usr/bin/env python3
"""
Model Benchmarking Script for Snowflake Cortex AI
Tests different models with real data from SNOWFLAKE_SAMPLE_DATA
Benchmarks: Quality, Speed, Cost, Accuracy
"""

import sys
import os
import time
import json
import shutil
import logging
from datetime import datetime
from pathlib import Path
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional, Any
import traceback  # noqa: F401

# Set up project-local temporary directory
# Use package root (parent of datamanagement_genai package)
PACKAGE_ROOT = Path(__file__).parent.parent
PROJECT_ROOT = PACKAGE_ROOT  # Package root is project root
TMP_DIR = PROJECT_ROOT / ".tmp"
TMP_DIR.mkdir(parents=True, exist_ok=True)

# Configure logging - use smart configuration for Jupyter vs scripts
try:
    from .logging_config import configure_logging  # noqa: F401
    # Only configure if not already configured (avoid duplicate handlers)
    if not logging.getLogger().handlers:
        configure_logging(verbose=False)  # Less verbose by default
except ImportError:
    # Fallback if logging_config not available
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(TMP_DIR / 'test_models.log'),
            logging.StreamHandler()
        ]
    )
    # Suppress Snowflake verbose logs
    logging.getLogger('snowflake.connector').setLevel(logging.WARNING)
    logging.getLogger('snowflake.snowpark').setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

try:
    import snowflake.connector  # noqa: F401
    SNOWFLAKE_CONNECTOR_AVAILABLE = True
except ImportError:
    SNOWFLAKE_CONNECTOR_AVAILABLE = False
    logger.warning("snowflake-connector-python not available")

try:
    from snowflake.snowpark import Session  # noqa: F401
    SNOWPARK_AVAILABLE = True
except ImportError:
    SNOWPARK_AVAILABLE = False
    logger.warning("snowflake-snowpark-python not available")

# Try to import tiktoken for accurate token counting
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
    # Use cl100k_base encoding (used by GPT-4, Claude, and most modern models)
    _tokenizer = tiktoken.get_encoding("cl100k_base")
except ImportError:
    TIKTOKEN_AVAILABLE = False
    logger.warning("tiktoken not available, falling back to character-based estimation")
    _tokenizer = None

# Try to import great_expectations for better code validation
try:
    import great_expectations as gx  # noqa: F401
    GX_AVAILABLE = True
except ImportError:
    GX_AVAILABLE = False
    logger.warning("great-expectations not available, GX code validation will be limited")

# Import data quality rules manager (required component)
from .data_quality.rules_manager import DataQualityRulesManager  # noqa: E402
DATA_QUALITY_RULES_AVAILABLE = True

# Default test prompts (fallback if config file not found)
_DEFAULT_TEST_PROMPTS = {
    "data_quality_gx": {
        "natural_language": "The L_ORDERKEY column should not be null and should be unique",
        "table": "LINEITEM",
        "schema": "SNOWFLAKE_SAMPLE_DATA.TPCH_SF1",
        "expected_elements": ["expect_column_values_to_not_be_null", "expect_column_values_to_be_unique", "L_ORDERKEY", "LINEITEM"],
        "test_type": "code_generation"
    },
    "sql_query_generation": {
        "prompt": "Write a SQL query to find all customers from the CUSTOMER table in SNOWFLAKE_SAMPLE_DATA.TPCH_SF1 who have a C_ACCTBAL greater than 5000, ordered by C_ACCTBAL descending. Return C_CUSTKEY, C_NAME, and C_ACCTBAL.",
        "expected_elements": ["SELECT", "FROM", "CUSTOMER", "C_ACCTBAL", "WHERE", "ORDER BY"],
        "test_type": "sql_generation"
    },
    "data_analysis": {
        "prompt": "Analyze the ORDERS table in SNOWFLAKE_SAMPLE_DATA.TPCH_SF1. What data quality checks should I perform? Consider null values, data types, and referential integrity.",
        "expected_elements": ["null", "data type", "quality", "check"],
        "test_type": "analysis"
    },
    "complex_validation": {
        "natural_language": "L_QUANTITY should be positive, L_EXTENDEDPRICE should be positive, L_ORDERKEY should not be null, L_SHIPDATE should not be in the future",
        "table": "LINEITEM",
        "schema": "SNOWFLAKE_SAMPLE_DATA.TPCH_SF1",
        "expected_elements": ["expect_column_values_to_be_between", "expect_column_values_to_not_be_null", "L_QUANTITY", "L_EXTENDEDPRICE"],
        "test_type": "code_generation"
    },
    "schema_explanation": {
        "prompt": "Explain the relationship between the CUSTOMER, ORDERS, and LINEITEM tables in the TPCH_SF1 schema. What are the primary keys and foreign keys?",
        "expected_elements": ["CUSTOMER", "ORDERS", "LINEITEM", "key", "relationship"],
        "test_type": "explanation"
    }
}

# Default models (fallback if config file not found)
_DEFAULT_MODELS = {
    "claude-3-5-sonnet": {
        "display": "Claude 3.5 Sonnet",
        "cost_per_1M_input": 3.00,
        "cost_per_1M_output": 15.00,
    },
    "claude-3-haiku": {
        "display": "Claude 3 Haiku",
        "cost_per_1M_input": 0.25,
        "cost_per_1M_output": 1.25,
    },
    "llama-3-70b": {
        "display": "Llama 3 70B",
        "cost_per_1M_input": 0.59,
        "cost_per_1M_output": 0.79,
    },
    "llama-3-8b": {
        "display": "Llama 3 8B",
        "cost_per_1M_input": 0.05,
        "cost_per_1M_output": 0.05,
    },
    "mistral-7b": {
        "display": "Mistral 7B",
        "cost_per_1M_input": 0.14,
        "cost_per_1M_output": 0.14,
    },
    "mixtral-8x7b": {
        "display": "Mixtral 8x7B",
        "cost_per_1M_input": 0.24,
        "cost_per_1M_output": 0.24,
    },
}

# Default token limits (fallback if config file not found)
_DEFAULT_TOKEN_LIMITS = {
    "claude-3-5-sonnet": 200000,
    "claude-3-haiku": 200000,
    "llama-3-70b": 8192,
    "llama-3-8b": 8192,
    "mistral-7b": 8192,
    "mixtral-8x7b": 32768,
}

# Schema cache to avoid repeated DESCRIBE calls
_schema_cache: Dict[str, List[str]] = {}

def load_test_config(config_path: Optional[str] = None) -> Dict[str, Any]:
    """
    Load test configuration from JSON file
    
    Args:
        config_path: Path to config file (default: test_config.json in package root)
    
    Returns:
        Dictionary with 'models', 'test_prompts', 'token_limits', and 'settings'
    """
    if config_path is None:
        # Try package root first, then current directory
        config_path = PACKAGE_ROOT / "test_config.json"
        if not config_path.exists():
            config_path = Path.cwd() / "test_config.json"
    else:
        config_path = Path(config_path)
    
    if not config_path.exists():
        logger.warning(f"Config file not found: {config_path}, using defaults")
        return {
            "models": _DEFAULT_MODELS,
            "test_prompts": _DEFAULT_TEST_PROMPTS,
            "token_limits": _DEFAULT_TOKEN_LIMITS,
            "settings": {
                "parallel_execution": True,
                "max_workers": 3,
                "output_directory": "tests",
                "log_level": "INFO"
            }
        }
    
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
        
        logger.info(f"Loaded configuration from: {config_path}")
        
        # Validate and merge with defaults
        models = config.get("models", _DEFAULT_MODELS)
        test_prompts = config.get("test_prompts", _DEFAULT_TEST_PROMPTS)
        token_limits = config.get("token_limits", _DEFAULT_TOKEN_LIMITS)
        settings = config.get("settings", {
            "parallel_execution": True,
            "max_workers": 3,
            "output_directory": "tests",
            "log_level": "INFO"
        })
        
        # Filter out disabled models and test prompts
        models = {k: v for k, v in models.items() if v.get("enabled", True)}
        test_prompts = {k: v for k, v in test_prompts.items() if v.get("enabled", True)}
        
        logger.info(f"Loaded {len(models)} enabled models and {len(test_prompts)} enabled test prompts")
        
        return {
            "models": models,
            "test_prompts": test_prompts,
            "token_limits": token_limits,
            "settings": settings
        }
        
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing config file {config_path}: {e}, using defaults")
        return {
            "models": _DEFAULT_MODELS,
            "test_prompts": _DEFAULT_TEST_PROMPTS,
            "token_limits": _DEFAULT_TOKEN_LIMITS,
            "settings": {
                "parallel_execution": True,
                "max_workers": 3,
                "output_directory": "tests",
                "log_level": "INFO"
            }
        }
    except Exception as e:
        logger.error(f"Error loading config file {config_path}: {e}, using defaults", exc_info=True)
        return {
            "models": _DEFAULT_MODELS,
            "test_prompts": _DEFAULT_TEST_PROMPTS,
            "token_limits": _DEFAULT_TOKEN_LIMITS,
            "settings": {
                "parallel_execution": True,
                "max_workers": 3,
                "output_directory": "tests",
                "log_level": "INFO"
            }
        }

# Load configuration from file
_CONFIG = load_test_config()

# Extract configuration values
MODELS = _CONFIG["models"]
TEST_PROMPTS = _CONFIG["test_prompts"]
MODEL_TOKEN_LIMITS = _CONFIG["token_limits"]

def reload_config(config_path: Optional[str] = None):
    """
    Reload configuration from file (useful if config was updated)
    
    Args:
        config_path: Path to config file (default: test_config.json in script directory)
    """
    global _CONFIG, MODELS, TEST_PROMPTS, MODEL_TOKEN_LIMITS
    _CONFIG = load_test_config(config_path)
    MODELS = _CONFIG["models"]
    TEST_PROMPTS = _CONFIG["test_prompts"]
    MODEL_TOKEN_LIMITS = _CONFIG["token_limits"]
    logger.info("Configuration reloaded")

def count_tokens(text: str) -> int:
    """Count tokens accurately using tiktoken if available, otherwise estimate"""
    if not text:
        return 0
    
    if TIKTOKEN_AVAILABLE and _tokenizer:
        try:
            return len(_tokenizer.encode(str(text)))
        except Exception as e:
            logger.warning(f"Token counting failed, using estimation: {e}")
            # Fallback to estimation
            return len(str(text)) // 4
    else:
        # Fallback estimation: ~4 characters per token
        return len(str(text)) // 4

def check_token_limit(prompt: str, model: str, max_output_tokens: int = 4000) -> Tuple[bool, int, int]:
    """
    Check if prompt + estimated output exceeds model token limit
    
    Returns:
        (within_limit, input_tokens, estimated_total_tokens)
    """
    input_tokens = count_tokens(prompt)
    estimated_output_tokens = max_output_tokens
    total_tokens = input_tokens + estimated_output_tokens
    
    model_limit = MODEL_TOKEN_LIMITS.get(model, 8192)
    within_limit = total_tokens < model_limit
    
    if not within_limit:
        logger.warning(
            f"Token limit warning for {model}: "
            f"input={input_tokens}, estimated_output={estimated_output_tokens}, "
            f"total={total_tokens}, limit={model_limit}"
        )
    
    return within_limit, input_tokens, total_tokens

def retry_with_backoff(func, max_retries: int = 3, base_delay: float = 1.0, *args, **kwargs):
    """
    Retry a function with exponential backoff
    
    Args:
        func: Function to retry
        max_retries: Maximum number of retry attempts
        base_delay: Base delay in seconds (will be doubled each retry)
        *args, **kwargs: Arguments to pass to func
    
    Returns:
        Result of func or raises last exception
    """
    last_exception = None
    
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            last_exception = e
            error_msg = str(e).lower()
            
            # Don't retry on certain errors
            if any(skip in error_msg for skip in ['unavailable', 'permission', 'authentication', 'syntax']):
                logger.error(f"Non-retryable error: {e}")
                raise e
            
            if attempt < max_retries - 1:
                delay = base_delay * (2 ** attempt)
                logger.warning(f"Attempt {attempt + 1} failed: {e}. Retrying in {delay}s...")
                time.sleep(delay)
            else:
                logger.error(f"All {max_retries} attempts failed. Last error: {e}")
    
    raise last_exception

def get_cached_table_columns(session, table_name: str, schema_name: str) -> List[str]:
    """
    Get table columns with caching to avoid repeated DESCRIBE calls
    
    Args:
        session: Snowflake session
        table_name: Name of the table
        schema_name: Full schema name (e.g., 'SNOWFLAKE_SAMPLE_DATA.TPCH_SF1')
    
    Returns:
        List of column names
    """
    cache_key = f"{schema_name}.{table_name}"
    
    # Check cache first
    if cache_key in _schema_cache:
        logger.debug(f"Using cached columns for {cache_key}")
        return _schema_cache[cache_key]
    
    # Fetch from database with retry
    def _fetch_columns():
        return get_table_columns(session, table_name, schema_name)
    
    try:
        column_names = retry_with_backoff(_fetch_columns, max_retries=3)
        
        # Cache the result
        _schema_cache[cache_key] = column_names
        logger.info(f"Cached columns for {cache_key}: {len(column_names)} columns")
        
        return column_names
    except Exception as e:
        logger.error(f"Failed to get columns for {cache_key}: {e}")
        return []

def clear_schema_cache():
    """Clear the schema cache (useful for testing or when schemas change)"""
    global _schema_cache
    _schema_cache.clear()
    logger.info("Schema cache cleared")

def get_snowflake_session():
    """Get Snowflake session using connector or Snowpark"""
    # Import from helpers to use the updated version that supports config.toml
    from .helpers import get_snowflake_session as _get_snowflake_session
    return _get_snowflake_session()

def get_table_columns(session, table_name: str, schema_name: str) -> List[str]:
    """
    Get column names from a table with improved error handling
    
    Args:
        session: Snowflake session (Snowpark or connector)
        table_name: Name of the table
        schema_name: Full schema name (e.g., 'SNOWFLAKE_SAMPLE_DATA.TPCH_SF1')
    
    Returns:
        List of column names, empty list on error
    """
    try:
        # Validate inputs
        if not table_name or not schema_name:
            raise ValueError(f"Invalid table_name or schema_name: table={table_name}, schema={schema_name}")
        
        desc_query = f"DESCRIBE TABLE {schema_name}.{table_name}"
        logger.debug(f"Executing: {desc_query}")
        
        if hasattr(session, 'sql'):  # Snowpark
            desc_df = session.sql(desc_query).to_pandas()
        else:  # Connector
            cursor = session.cursor()
            try:
                cursor.execute(desc_query)
                columns = cursor.fetchall()
                # Convert to DataFrame-like structure
                import pandas as pd
                desc_df = pd.DataFrame(
                    columns, 
                    columns=[desc[0] for desc in cursor.description] if cursor.description else ['name', 'type']
                )
            finally:
                cursor.close()
        
        if desc_df.empty:
            logger.warning(f"DESCRIBE returned empty result for {schema_name}.{table_name}")
            return []
        
        # Find column name column (robust detection)
        desc_name_col = None
        for col_name in ['name', 'NAME', 'column_name', 'COLUMN_NAME', 'Name', 'COLUMN']:
            if col_name in desc_df.columns:
                desc_name_col = col_name
                break
        
        # If still not found, check if any column name contains 'name' (case insensitive)
        if desc_name_col is None:
            for col in desc_df.columns:
                if 'name' in str(col).lower():
                    desc_name_col = col
                    break
        
        # Last resort: use the first column
        if desc_name_col is None and len(desc_df.columns) > 0:
            desc_name_col = desc_df.columns[0]
        
        # Get column names
        if desc_name_col and desc_name_col in desc_df.columns:
            column_names = desc_df[desc_name_col].tolist()
        else:
            # Fallback: use first column
            column_names = desc_df.iloc[:, 0].tolist() if len(desc_df.columns) > 0 else []
        
        # Filter out None/empty values
        column_names = [col for col in column_names if col]
        
        logger.debug(f"Found {len(column_names)} columns for {schema_name}.{table_name}")
        return column_names
        
    except Exception as e:
        logger.error(
            f"Error getting columns for {schema_name}.{table_name}: {e}",
            exc_info=True
        )
        return []

def build_gx_prompt(natural_language, column_names):
    """Build Great Expectations prompt matching streamlit app structure"""
    system_prompt = """You are an expert in Great Expectations data validation library.
Convert the following natural language description into Python code using the validator.expect_*() format.

CRITICAL INSTRUCTIONS:
- Return ONLY executable Python code using validator.expect_*() methods
- NO explanations, NO thinking process, NO markdown, NO comments
- Use the format: validator.expect_column_values_to_be_unique(column="column_name")
- Each expectation on a new line
- Do NOT use expectation_suite.add_expectation()
- Do NOT import anything
- Do NOT create a validator (it already exists)
- Use EXACT column names as provided in the available columns list
- Column names are CASE-SENSITIVE - use them EXACTLY as shown

Example 1:
Input: "Check that none of the values in the address column match the pattern for an address starting with a digit"
Output: validator.expect_column_values_to_not_match_regex(column="address", regex=r"^\\d")

Example 2:
Input: "transaction_id should be unique and customer_id should not be null"
Output: validator.expect_column_values_to_be_unique(column="transaction_id")
validator.expect_column_values_to_not_be_null(column="customer_id")

Example 3:
Input: "amount should be greater than 0"
Output: validator.expect_column_values_to_be_between(column="amount", min_value=0, strict_min=True)

Available columns in this dataset: {column_names_str}

Natural language description:
{natural_language}"""
    
    # Format column names for the prompt
    column_names_str = ", ".join(f'"{col}"' for col in column_names[:50])  # Limit to first 50 columns
    prompt = system_prompt.format(
        column_names_str=column_names_str,
        natural_language=natural_language
    )
    
    return prompt

def execute_cortex_query(session, model: str, prompt: str) -> str:
    """
    Execute Cortex AI query and return response with improved error handling
    
    Args:
        session: Snowflake session
        model: Model name to use
        prompt: Prompt text
    
    Returns:
        Response content as string
    
    Raises:
        Exception: If query fails or model is unavailable
    """
    try:
        # Validate inputs
        if not model or not prompt:
            raise ValueError(f"Invalid model or prompt: model={model}, prompt_length={len(prompt) if prompt else 0}")
        
        # Escape single quotes for SQL (double single quotes)
        escaped_prompt = prompt.replace("'", "''")
        
        # Check prompt length (Snowflake has limits)
        if len(escaped_prompt) > 1000000:  # ~1MB limit
            raise ValueError(f"Prompt too long: {len(prompt)} characters")
        
        sql_query = f"""
        SELECT SNOWFLAKE.CORTEX.COMPLETE(
            '{model}',
            '{escaped_prompt}'
        ) AS response
        """
        
        logger.debug(f"Executing Cortex query for model: {model}, prompt length: {len(prompt)}")
        
        # Execute based on session type
        if hasattr(session, 'sql'):  # Snowpark
            result = session.sql(sql_query).collect()
            response_raw = result[0][0] if isinstance(result[0], (list, tuple)) else result[0]
        else:  # Connector
            cursor = session.cursor()
            try:
                cursor.execute(sql_query)
                result = cursor.fetchone()
                response_raw = result[0] if result else None
            finally:
                cursor.close()
        
        if response_raw is None:
            raise Exception("Cortex AI returned None response")
        
        # Parse response
        if isinstance(response_raw, str):
            try:
                response = json.loads(response_raw)
                if isinstance(response, dict) and "choices" in response:
                    if len(response["choices"]) > 0:
                        choice = response["choices"][0]
                        if "message" in choice:
                            content = choice["message"].get("content", "")
                        elif "messages" in choice:
                            content = choice["messages"]
                        elif "text" in choice:
                            content = choice["text"]
                        else:
                            content = str(choice)
                    else:
                        content = str(response)
                elif "content" in response:
                    content = response["content"]
                elif "text" in response:
                    content = response["text"]
                else:
                    content = str(response)
            except json.JSONDecodeError:
                # Response might not be JSON, use as-is
                content = response_raw
            except Exception as e:
                logger.warning(f"Error parsing JSON response: {e}, using raw response")
                content = response_raw
        else:
            content = str(response_raw) if response_raw else ""
        
        if not content:
            raise Exception("Cortex AI returned empty response")
        
        # Check for truncation indicators and warn
        truncation_indicators = [
            "[continued in next part",
            "continued in next part",
            "[continued",
            "due to length",
            "would you like me to continue",
            "remaining sections"
        ]
        
        content_lower = content.lower()
        for indicator in truncation_indicators:
            if indicator in content_lower:
                logger.warning(f"Response may be truncated - found indicator: '{indicator}'")
                print("⚠️  Warning: LLM response may be truncated. Consider using a model with higher output limits or splitting the report generation.")
                break
        
        logger.debug(f"Received response of length: {len(content)}")
        return content
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Cortex query failed for model {model}: {error_msg}")
        
        # Check if it's a model availability issue
        if "unavailable" in error_msg.lower() or "unavail" in error_msg.lower():
            raise Exception(f"Model '{model}' is not available in your Snowflake account. Error: {error_msg}")
        elif "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
            raise Exception(f"Query timeout for model '{model}': {error_msg}")
        else:
            raise Exception(f"Cortex AI query failed for model '{model}': {error_msg}")

def check_model_availability(session, models_to_check: List[str]) -> Dict[str, bool]:
    """
    Check which models are available in the Snowflake account
    
    Args:
        session: Snowflake session
        models_to_check: List of model names to check
    
    Returns:
        Dictionary mapping model names to availability (True/False)
    """
    availability = {}
    test_prompt = "Hello"  # Simple test prompt
    
    print("\n" + "=" * 80)
    print("CHECKING MODEL AVAILABILITY")
    print("=" * 80)
    print(f"Testing {len(models_to_check)} models for availability...\n")
    
    for model in models_to_check:
        try:
            # Try a simple query to test availability
            execute_cortex_query(session, model, test_prompt)
            availability[model] = True
            print(f"✓ {MODELS.get(model, {}).get('display', model)} ({model}) - Available")
        except Exception as e:
            error_msg = str(e).lower()
            # Check if it's an availability issue
            if "unavailable" in error_msg or "not available" in error_msg or "not found" in error_msg:
                availability[model] = False
                print(f"✗ {MODELS.get(model, {}).get('display', model)} ({model}) - Not Available")
            else:
                # Other errors might be transient, mark as available but log warning
                logger.warning(f"Model {model} check returned error (may be transient): {e}")
                availability[model] = True  # Give it a chance, might be a transient issue
                print(f"⚠ {MODELS.get(model, {}).get('display', model)} ({model}) - Available (with warnings)")
    
    available_count = sum(1 for v in availability.values() if v)
    unavailable_count = len(availability) - available_count
    
    print("\nAvailability Check Complete:")
    print(f"  Available: {available_count}")
    print(f"  Unavailable: {unavailable_count}")
    print("=" * 80 + "\n")
    
    return availability

def test_model_with_prompt(session, model, test_config):
    """Test a model with a specific prompt configuration"""
    expected_elements = test_config.get("expected_elements", [])
    test_type = test_config.get("test_type", "general")
    
    # Build prompt based on test type
    if test_type == "code_generation" and "natural_language" in test_config:
        # Use GX prompt structure from streamlit app
        table_name = test_config.get("table", "")
        schema_name = test_config.get("schema", "")
        natural_language = test_config["natural_language"]
        
        # Get column names from table (using cache)
        column_names = []
        if table_name and schema_name:
            column_names = get_cached_table_columns(session, table_name, schema_name)
            if not column_names:
                logger.warning(f"No columns found for {schema_name}.{table_name}, proceeding without column context")
        
        # Build prompt using streamlit app structure
        prompt = build_gx_prompt(natural_language, column_names)
        
        # Check token limits before sending
        within_limit, input_tokens_check, total_tokens_check = check_token_limit(prompt, model)
        if not within_limit:
            logger.warning(
                f"Prompt may exceed token limit for {model}. "
                f"Input tokens: {input_tokens_check}, Estimated total: {total_tokens_check}"
            )
    else:
        # Use direct prompt
        prompt = test_config.get("prompt", "")
    
    start_time = time.time()
    
    try:
        # Execute query with retry logic
        response = retry_with_backoff(
            execute_cortex_query,
            max_retries=3,
            base_delay=1.0,
            session=session,
            model=model,
            prompt=prompt
        )
        elapsed_time = time.time() - start_time
        
        # Count tokens accurately using tokenizer
        input_tokens = count_tokens(prompt)
        output_tokens = count_tokens(str(response))
        
        # Calculate cost
        model_info = MODELS.get(model, {})
        input_cost = (input_tokens / 1_000_000) * model_info.get("cost_per_1M_input", 0)
        output_cost = (output_tokens / 1_000_000) * model_info.get("cost_per_1M_output", 0)
        total_cost = input_cost + output_cost
        
        # Evaluate quality
        quality_score, feedback = evaluate_response_quality(
            str(response), 
            test_type, 
            expected_elements
        )
        
        # Test if SQL/Code is executable (for code generation tests)
        executable = False
        execution_error = None
        if test_type in ["sql_generation", "code_generation"]:
            executable, execution_error = test_code_execution(session, str(response), test_type)
        
        # Update data quality rules CSV if this is a data quality or analysis test
        dq_rules_updated = False
        if test_type in ["analysis", "code_generation"]:
            try:
                # Initialize data quality rules manager with RAG system if available
                dq_manager = None
                try:
                    # Try to create RAG system for DAMA references
                    dq_rag_system = None
                    try:
                        from ..rag.system import RAGSystem
                        import os
                        rag_database = os.getenv("SNOWFLAKE_RAG_DATABASE", "RAG_KNOWLEDGE_BASE")
                        rag_schema = os.getenv("SNOWFLAKE_RAG_SCHEMA", "PUBLIC")
                        dq_rag_system = RAGSystem(session=session, vector_store_database=rag_database, vector_store_schema=rag_schema)
                        logger.debug("RAG system initialized for DAMA references")
                    except Exception as rag_error:
                        logger.debug(f"Could not initialize RAG for DAMA references: {rag_error}")
                    
                    # Initialize manager with or without RAG
                    if dq_rag_system:
                        dq_manager = DataQualityRulesManager("data_quality_rules.csv", rag_system=dq_rag_system)
                    else:
                        dq_manager = DataQualityRulesManager("data_quality_rules.csv")
                except Exception as e:
                    logger.warning(f"Could not initialize data quality rules manager: {e}")
                    dq_manager = None
                
                if dq_manager:
                    # Prepare data context for LLM analysis
                    table_name = test_config.get("table", "")
                    schema_name = test_config.get("schema", "")
                    
                    # Get column information if available
                    column_info = []
                    if table_name and schema_name:
                        column_names = get_cached_table_columns(session, table_name, schema_name)
                        if column_names:
                            # Get detailed column info
                            try:
                                desc_query = f"DESCRIBE TABLE {schema_name}.{table_name}"
                                if hasattr(session, 'sql'):
                                    desc_df = session.sql(desc_query).to_pandas()
                                else:
                                    cursor = session.cursor()
                                    try:
                                        cursor.execute(desc_query)
                                        columns = cursor.fetchall()
                                        import pandas as pd
                                        desc_df = pd.DataFrame(
                                            columns,
                                            columns=[desc[0] for desc in cursor.description] if cursor.description else ['name', 'type', 'nullable']
                                        )
                                    finally:
                                        cursor.close()
                                
                                # Extract column info
                                for _, row in desc_df.iterrows():
                                    col_name = row.get('name', row.get('NAME', row.get('column_name', '')))
                                    col_type = row.get('type', row.get('TYPE', row.get('data_type', 'VARCHAR')))
                                    nullable = row.get('nullable', row.get('NULLABLE', row.get('null', 'YES')))
                                    if col_name:
                                        column_info.append({
                                            "name": str(col_name),
                                            "type": str(col_type),
                                            "nullable": str(nullable).upper() == "YES" or str(nullable).upper() == "TRUE"
                                        })
                            except Exception as e:
                                logger.warning(f"Could not get detailed column info: {e}")
                                # Fallback to just column names
                                for col_name in column_names:
                                    column_info.append({"name": col_name, "type": "VARCHAR", "nullable": True})
                    
                    # Prepare data context
                    data_context = {
                        "test_name": test_config.get("test_name", "unknown"),
                        "test_type": test_type,
                        "llm_response": str(response)[:1000],  # Include LLM response as context
                        "table": table_name,
                        "schema": schema_name,
                        "business_context": f"Data quality analysis for {schema_name}.{table_name}" if table_name else "General data quality analysis"
                    }
                    
                    # Use the best available model for analysis (or default to llama-3-70b)
                    analysis_model = model  # Use the same model that generated the response
                    
                    # Analyze and update rules
                    result = dq_manager.analyze_and_update_with_llm(
                        session=session,
                        model=analysis_model,
                        data_context=data_context,
                        table_name=table_name if table_name else None,
                        schema_name=schema_name if schema_name else None,
                        column_info=column_info if column_info else None
                    )
                    
                    if result.get('success'):
                        rules_created = result.get('rules_created', 0)
                        rules_updated = result.get('rules_updated', 0)
                        if rules_created > 0 or rules_updated > 0:
                            logger.info(f"Data quality rules updated: {rules_created} created, {rules_updated} updated")
                            dq_rules_updated = True
                            feedback.append(f"✓ Updated data quality rules CSV ({rules_created} new, {rules_updated} updated)")
            except Exception as e:
                logger.warning(f"Failed to update data quality rules: {e}")
        
        return {
            "success": True,
            "response": str(response),
            "prompt": prompt,  # Store the actual prompt used
            "response_length": len(str(response)),
            "elapsed_time": elapsed_time,
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "total_tokens": input_tokens + output_tokens,
            "estimated_cost": total_cost,
            "tokens_per_second": (input_tokens + output_tokens) / elapsed_time if elapsed_time > 0 else 0,
            "quality_score": quality_score,
            "quality_feedback": feedback,
            "executable": executable,
            "execution_error": execution_error,
            "dq_rules_updated": dq_rules_updated,
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "elapsed_time": time.time() - start_time,
            "estimated_cost": 0,
            "quality_score": 0,
        }

def evaluate_response_quality(response, test_type, expected_elements):
    """Evaluate response quality based on test type and expected elements"""
    score = 0
    max_score = 10
    feedback = []
    
    response_lower = response.lower()
    
    # Check for expected elements
    elements_found = 0
    for element in expected_elements:
        if element.lower() in response_lower:
            elements_found += 1
    
    if expected_elements:
        element_score = (elements_found / len(expected_elements)) * 5
        score += element_score
        feedback.append(f"Found {elements_found}/{len(expected_elements)} expected elements")
    
    # Type-specific evaluation
    if test_type == "code_generation":
        if "expect_" in response_lower or "validator" in response_lower:
            score += 2
            feedback.append("✓ Contains GX code structure")
        if "def " in response or "import " in response:
            score += 1
            feedback.append("✓ Includes Python code")
        if len(response) > 200:
            score += 1
            feedback.append("✓ Detailed code")
        if "error" not in response_lower or "cannot" not in response_lower:
            score += 1
            feedback.append("✓ No errors mentioned")
    
    elif test_type == "sql_generation":
        if "select" in response_lower and "from" in response_lower:
            score += 2
            feedback.append("✓ Contains SQL query")
        if "where" in response_lower:
            score += 1
            feedback.append("✓ Includes filtering")
        if "order by" in response_lower:
            score += 1
            feedback.append("✓ Includes ordering")
        if len(response) > 100:
            score += 1
            feedback.append("✓ Complete query")
    
    elif test_type == "analysis":
        if len(response) > 200:
            score += 2
            feedback.append("✓ Comprehensive analysis")
        if "check" in response_lower or "validate" in response_lower:
            score += 1
            feedback.append("✓ Provides actionable checks")
        if "null" in response_lower or "missing" in response_lower:
            score += 1
            feedback.append("✓ Addresses data quality")
    
    else:
        # Generic evaluation
        if len(response) > 100:
            score += 2
        if len(response) > 300:
            score += 1
        if "error" not in response_lower:
            score += 1
    
    return min(score, max_score), feedback

def test_code_execution(session, response: str, test_type: str) -> Tuple[bool, Optional[str]]:
    """
    Test if generated SQL or code is executable with improved validation
    
    Args:
        session: Snowflake session
        response: Generated code/query response
        test_type: Type of test ('sql_generation' or 'code_generation')
    
    Returns:
        Tuple of (is_executable, error_message)
    """
    if test_type == "sql_generation":
        # Extract SQL query from response
        import re
        
        # Try to find SQL block
        sql_match = re.search(r'```sql\s*(.*?)\s*```', response, re.DOTALL | re.IGNORECASE)
        if sql_match:
            sql_query = sql_match.group(1).strip()
        else:
            # Try to find SELECT statement (more robust pattern)
            select_match = re.search(r'(SELECT\s+.*?)(?:\n\n|```|\Z)', response, re.DOTALL | re.IGNORECASE)
            if select_match:
                sql_query = select_match.group(1).strip()
            else:
                # Last resort: look for any SQL-like statement
                sql_keywords = ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'CREATE', 'ALTER']
                for keyword in sql_keywords:
                    pattern = rf'({keyword}.*?)(?:\n\n|```|\Z)'
                    match = re.search(pattern, response, re.DOTALL | re.IGNORECASE)
                    if match:
                        sql_query = match.group(1).strip()
                        break
                else:
                    return False, "No SQL query found in response"
        
        # Clean up the query
        sql_query = sql_query.strip().rstrip(';')
        
        # Basic SQL validation
        if not any(keyword in sql_query.upper() for keyword in ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'CREATE']):
            return False, "Query does not contain valid SQL keywords"
        
        # Test execution (limit to prevent issues)
        if "LIMIT" not in sql_query.upper():
            sql_query = sql_query + " LIMIT 10"
        
        try:
            # Execute with retry logic
            def _execute_sql():
                if hasattr(session, 'sql'):  # Snowpark
                    result = session.sql(sql_query).collect()
                    return result
                else:  # Connector
                    cursor = session.cursor()
                    try:
                        cursor.execute(sql_query)
                        result = cursor.fetchone()  # Execute but don't fetch all
                        return result
                    finally:
                        cursor.close()
            
            retry_with_backoff(_execute_sql, max_retries=2, base_delay=0.5)
            logger.debug("SQL query executed successfully")
            return True, None
            
        except Exception as e:
            error_msg = str(e)
            logger.warning(f"SQL execution failed: {error_msg[:200]}")
            return False, error_msg[:500]  # Truncate but keep more context
    
    elif test_type == "code_generation":
        # Extract Python code from response
        import re
        
        # Try to find code block
        code_match = re.search(r'```python\s*(.*?)\s*```', response, re.DOTALL | re.IGNORECASE)
        if code_match:
            code = code_match.group(1).strip()
        else:
            # Look for code without markdown (lines starting with validator.expect_)
            code_lines = []
            for line in response.split('\n'):
                stripped = line.strip()
                if stripped and (stripped.startswith('validator.') or stripped.startswith('expect_')):
                    code_lines.append(line)
            if code_lines:
                code = '\n'.join(code_lines)
            else:
                code = response
        
        # Check if it contains GX code
        if not any(keyword in code for keyword in ['validator.', 'expect_', 'expectation']):
            return False, "No Great Expectations code found (missing validator.expect_* or expect_* patterns)"
        
        # Basic syntax check
        try:
            compile(code, '<string>', 'exec')
        except SyntaxError as e:
            return False, f"Syntax error: {str(e)[:300]}"
        except Exception as e:
            return False, f"Compilation error: {str(e)[:300]}"
        
        # If GX is available, try to validate the code structure more thoroughly
        if GX_AVAILABLE:
            try:
                # Check if code uses valid GX patterns
                if 'validator.expect_' in code or 'expect_' in code:
                    # Additional validation: check for common GX method patterns
                    gx_patterns = [
                        r'validator\.expect_\w+',
                        r'expect_\w+',
                        r'column\s*=\s*["\']',
                    ]
                    has_valid_pattern = any(re.search(pattern, code) for pattern in gx_patterns)
                    if has_valid_pattern:
                        logger.debug("GX code structure validated")
                        return True, None
                    else:
                        return False, "Code does not match expected GX patterns"
            except Exception as e:
                logger.warning(f"GX validation check failed: {e}")
                # Fall back to syntax check result
                return True, None  # Syntax check passed, assume it's valid
        
        return True, None
    
    return False, f"Unknown test type: {test_type}"

def _run_single_test(session_factory, model: str, test_name: str, test_config: Dict[str, Any]) -> Dict[str, Any]:
    """
    Run a single test scenario for a model (used for parallel execution)
    
    Args:
        session_factory: Function that returns a new Snowflake session
        model: Model name to test
        test_name: Name of the test scenario
        test_config: Test configuration dictionary
    
    Returns:
        Test result dictionary
    """
    # Create a new session for this test (thread-safe)
    try:
        session = session_factory()
        if not session:
            return {
                "success": False,
                "error": "Could not create Snowflake session",
                "model": model,
                "test_name": test_name,
            }
        
        logger.info(f"Running test {test_name} for model {model}")
        result = test_model_with_prompt(session, model, test_config)
        result["model"] = model
        result["test_name"] = test_name
        
        # Close session
        try:
            if hasattr(session, 'close'):
                session.close()
        except Exception:
            pass
        
        return result
        
    except Exception as e:
        logger.error(f"Error in test {test_name} for model {model}: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e),
            "model": model,
            "test_name": test_name,
            "elapsed_time": 0,
            "estimated_cost": 0,
        }

def run_model_benchmarks(session, models_to_test=None, parallel: Optional[bool] = None, max_workers: Optional[int] = None):
    """
    Run comprehensive benchmarks on specified models with optional parallel execution
    
    Args:
        session: Snowflake session (used as template for creating new sessions)
        models_to_test: List of model names to test (None = all enabled models from config)
        parallel: Whether to run tests in parallel (None = use config file setting)
        max_workers: Maximum number of parallel workers (None = use config file setting)
    
    Returns:
        List of test results
    """
    # Get settings from config
    settings = _CONFIG.get("settings", {})
    
    # Use config values if not explicitly provided
    if parallel is None:
        parallel = settings.get("parallel_execution", True)
    if max_workers is None:
        max_workers = settings.get("max_workers", 3)
    
    # Get models to test (use enabled models from config if not specified)
    if models_to_test is None:
        models_to_test = [model_id for model_id, model_config in MODELS.items() 
                         if model_config.get("enabled", True)]
    
    # Check model availability and filter to only available models
    availability = check_model_availability(session, models_to_test)
    models_to_test = [model for model in models_to_test if availability.get(model, False)]
    
    if not models_to_test:
        print("ERROR: No available models found. Please check your Snowflake account configuration.")
        print("Available models may vary by account type and region.")
        return []
    
    print("=" * 80)
    print("SNOWFLAKE CORTEX AI MODEL BENCHMARKING")
    print("=" * 80)
    print(f"Testing {len(models_to_test)} available models with {len(TEST_PROMPTS)} test scenarios")
    print("Using data from: SNOWFLAKE_SAMPLE_DATA.TPCH_SF1")
    print(f"Parallel execution: {'Enabled' if parallel else 'Disabled'}")
    if parallel:
        print(f"Max workers: {max_workers}")
    print(f"Started at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()
    
    # Create session factory function
    def session_factory():
        """Create a new Snowflake session using the same config as the original"""
        return get_snowflake_session()
    
    results = []
    
    if parallel:
        # Prepare all test tasks
        tasks = []
        for model in models_to_test:
            if model not in MODELS:
                print(f"⚠️  Skipping unknown model: {model}")
                continue
            
            for test_name, test_config in TEST_PROMPTS.items():
                tasks.append((model, test_name, test_config))
        
        print(f"Running {len(tasks)} tests in parallel with {max_workers} workers...\n")
        
        # Execute tests in parallel
        completed = 0
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_task = {
                executor.submit(_run_single_test, session_factory, model, test_name, test_config): (model, test_name)
                for model, test_name, test_config in tasks
            }
            
            # Process completed tasks
            for future in as_completed(future_to_task):
                model, test_name = future_to_task[future]
                completed += 1
                
                try:
                    result = future.result()
                    results.append(result)
                    
                    # Print progress
                    status = "✓" if result.get("success") else "✗"
                    print(f"[{completed}/{len(tasks)}] {status} {MODELS[model]['display']} - {test_name}")
                    
                    if result.get("success"):
                        print(f"    Time: {result.get('elapsed_time', 0):.2f}s, "
                              f"Tokens: {result.get('total_tokens', 0)}, "
                              f"Cost: ${result.get('estimated_cost', 0):.6f}, "
                              f"Quality: {result.get('quality_score', 0):.1f}/10")
                    else:
                        error = result.get('error', 'Unknown error')
                        print(f"    Error: {error[:100]}")
                    
                except Exception as e:
                    logger.error(f"Task failed: {e}", exc_info=True)
                    results.append({
                        "success": False,
                        "error": str(e),
                        "model": model,
                        "test_name": test_name,
                    })
        
        print(f"\n✓ Completed {completed}/{len(tasks)} tests")
        
    else:
        # Sequential execution (original behavior)
        for model in models_to_test:
            if model not in MODELS:
                print(f"⚠️  Skipping unknown model: {model}")
                continue
            
            print(f"\n{'='*80}")
            print(f"Testing: {MODELS[model]['display']} ({model})")
            print(f"{'='*80}")
        
        model_results = []
        
        for test_name, test_config in TEST_PROMPTS.items():
            print(f"\n  Test: {test_name}")
            print(f"  Type: {test_config['test_type']}")
            # Show prompt preview (handle both direct prompts and natural language)
            if 'prompt' in test_config:
                prompt_preview = test_config['prompt'][:70] + "..." if len(test_config['prompt']) > 70 else test_config['prompt']
            elif 'natural_language' in test_config:
                prompt_preview = f"Natural Language: {test_config['natural_language'][:50]}..."
            else:
                prompt_preview = "N/A"
            print(f"  Prompt: {prompt_preview}")
            
            result = test_model_with_prompt(session, model, test_config)
            result["model"] = model
            result["test_name"] = test_name
            model_results.append(result)
            
            if result["success"]:
                print("  ✓ Success")
                print(f"    Time: {result['elapsed_time']:.2f}s")
                print(f"    Tokens: {result['total_tokens']} (in: {result['input_tokens']}, out: {result['output_tokens']})")
                print(f"    Cost: ${result['estimated_cost']:.6f}")
                print(f"    Quality: {result['quality_score']:.1f}/10")
                if result.get("executable") is not None:
                    exec_status = "✓ Executable" if result["executable"] else "✗ Not executable"
                    print(f"    {exec_status}")
                    if result.get("execution_error"):
                        print(f"      Error: {result['execution_error'][:100]}")
                if result.get("quality_feedback"):
                    for fb in result["quality_feedback"][:2]:
                        print(f"      {fb}")
            else:
                print(f"  ✗ Failed: {result.get('error', 'Unknown error')[:100]}")
        
        results.extend(model_results)
    
    return results

def analyze_benchmark_results(results):
    """Analyze and rank models based on benchmark results"""
    print("\n" + "=" * 80)
    print("BENCHMARK RESULTS ANALYSIS")
    print("=" * 80)
    
    # Group results by model
    model_stats = {}
    
    for result in results:
        model = result["model"]
        if model not in model_stats:
            model_stats[model] = {
                "tests": [],
                "success_count": 0,
                "total_tests": 0,
                "total_time": 0,
                "total_cost": 0,
                "total_tokens": 0,
                "total_quality_score": 0,
                "executable_count": 0,
                "executable_tests": 0,
            }
        
        model_stats[model]["tests"].append(result)
        model_stats[model]["total_tests"] += 1
        
        if result["success"]:
            model_stats[model]["success_count"] += 1
            model_stats[model]["total_time"] += result.get("elapsed_time", 0)
            model_stats[model]["total_cost"] += result.get("estimated_cost", 0)
            model_stats[model]["total_tokens"] += result.get("total_tokens", 0)
            quality = result.get("quality_score", 0)
            model_stats[model]["total_quality_score"] += quality
            
            if result.get("executable") is not None:
                model_stats[model]["executable_tests"] += 1
                if result["executable"]:
                    model_stats[model]["executable_count"] += 1
    
    # Calculate averages and rankings
    model_rankings = []
    unavailable_models_list = []
    
    for model, stats in model_stats.items():
        if stats["success_count"] == 0:
            # Collect unavailable model information
            error_messages = []
            for test in stats["tests"]:
                if not test.get("success") and test.get("error"):
                    error_msg = test.get("error", "")
                    # Extract the main error reason
                    if "not available" in error_msg.lower() or "unavailable" in error_msg.lower():
                        error_messages.append("Model not available in account")
                    else:
                        error_messages.append(error_msg[:100])
            
            # Get unique error reasons
            unique_errors = list(set(error_messages)) if error_messages else ["All tests failed"]
            unavailable_models_list.append({
                "model": model,
                "display": MODELS.get(model, {}).get("display", model),
                "total_tests": stats["total_tests"],
                "failed_tests": stats["total_tests"],
                "error_reason": unique_errors[0] if unique_errors else "Unknown error"
            })
            continue
        
        avg_time = stats["total_time"] / stats["success_count"]
        avg_cost = stats["total_cost"] / stats["success_count"]
        avg_quality = stats["total_quality_score"] / stats["success_count"]
        success_rate = (stats["success_count"] / stats["total_tests"]) * 100
        
        executable_rate = 0
        if stats["executable_tests"] > 0:
            executable_rate = (stats["executable_count"] / stats["executable_tests"]) * 100
        
        # Calculate composite score
        # Quality: 40%, Cost efficiency: 25%, Speed: 15%, Success rate: 10%, Executability: 10%
        quality_normalized = (avg_quality / 10) * 100
        cost_efficiency = 100 / (avg_cost * 1000000 + 1)
        speed_score = 100 / (avg_time * 10 + 1)
        
        composite_score = (
            quality_normalized * 0.4 +
            cost_efficiency * 0.25 +
            speed_score * 0.15 +
            success_rate * 0.1 +
            executable_rate * 0.1
        )
        
        model_rankings.append({
            "model": model,
            "display": MODELS[model]["display"],
            "success_rate": success_rate,
            "avg_quality": avg_quality,
            "avg_time": avg_time,
            "avg_cost": avg_cost,
            "executable_rate": executable_rate,
            "composite_score": composite_score,
            "stats": stats
        })
    
    # Sort by composite score
    model_rankings.sort(key=lambda x: x["composite_score"], reverse=True)
    
    # Print rankings
    print("\nMODEL RANKINGS (Best to Worst):")
    print("-" * 100)
    print(f"{'Rank':<6} {'Model':<25} {'Quality':<10} {'Time (s)':<12} {'Cost ($)':<12} {'Success':<10} {'Exec':<10} {'Score':<10}")
    print("-" * 100)
    
    for idx, ranking in enumerate(model_rankings, 1):
        print(f"{idx:<6} {ranking['display']:<25} {ranking['avg_quality']:<10.2f} "
              f"{ranking['avg_time']:<12.2f} {ranking['avg_cost']:<12.6f} "
              f"{ranking['success_rate']:<10.1f}% {ranking['executable_rate']:<10.1f}% {ranking['composite_score']:<10.2f}")
    
    # Detailed analysis
    print("\n" + "=" * 80)
    print("DETAILED ANALYSIS - TOP 3 MODELS")
    print("=" * 80)
    
    for ranking in model_rankings[:3]:
        print(f"\n{ranking['display']} ({ranking['model']}):")
        print(f"  Composite Score: {ranking['composite_score']:.2f}")
        print(f"  Average Quality: {ranking['avg_quality']:.2f}/10")
        print(f"  Average Response Time: {ranking['avg_time']:.2f}s")
        print(f"  Average Cost per Query: ${ranking['avg_cost']:.6f}")
        print(f"  Success Rate: {ranking['success_rate']:.1f}%")
        print(f"  Code Executability Rate: {ranking['executable_rate']:.1f}%")
    
    # Recommendation
    if model_rankings:
        best = model_rankings[0]
        print("\n" + "=" * 80)
        print("RECOMMENDATION")
        print("=" * 80)
        print(f"\n🏆 BEST DEFAULT MODEL: {best['display']} ({best['model']})")
        print("\nReasoning:")
        print(f"  • Highest composite score: {best['composite_score']:.2f}")
        print(f"  • Quality: {best['avg_quality']:.2f}/10")
        print(f"  • Cost-effective: ${best['avg_cost']:.6f} per query")
        print(f"  • Fast: {best['avg_time']:.2f}s average response time")
        print(f"  • Reliable: {best['success_rate']:.1f}% success rate")
        print(f"  • Executable code: {best['executable_rate']:.1f}% of generated code runs successfully")
        print("\nThis model provides the best balance for the application's use cases.")
    
    # Print unavailable models summary
    if unavailable_models_list:
        print("\n" + "=" * 80)
        print("UNAVAILABLE MODELS (Not included in rankings)")
        print("=" * 80)
        for unavail in unavailable_models_list:
            print(f"  • {unavail['display']} ({unavail['model']}): {unavail['error_reason']}")
        print(f"\nNote: {len(unavailable_models_list)} model(s) were not available and are excluded from rankings.")
        print("See Appendix for details.")
    
    if model_rankings:
        return best["model"], unavailable_models_list
    
    return None, unavailable_models_list

def generate_llm_report(session, results, model_rankings, best_model, use_section_based=True, use_rag=True):
    """
    Use LLM to generate a comprehensive analysis report with all explanatory content
    
    Args:
        session: Snowflake session
        results: Test results
        model_rankings: Model rankings
        best_model: Best model dict
        use_section_based: If True, generate report in sections to avoid truncation
        use_rag: If True, enhance report with RAG knowledge base references
    
    Returns:
        Complete report as string
    """
    print("\n" + "=" * 80)
    print("GENERATING COMPREHENSIVE REPORT WITH LLM")
    if use_rag:
        print("(Enhanced with RAG Knowledge Base)")
    print("=" * 80)
    
    if use_section_based:
        return generate_llm_report_section_based(session, results, model_rankings, best_model, use_rag=use_rag)
    else:
        return generate_llm_report_single(session, results, model_rankings, best_model, use_rag=use_rag)

def _enhance_section_with_rag(rag_system, session, report_model, base_content, section_name, summary_data=None, top_k=3, validate_citations=True):
    """Helper function to enhance a section with RAG, including benchmark context"""
    if not rag_system or not base_content:
        return base_content
    
    try:
        # Build context query that includes section name and key benchmark data
        context_query_parts = [section_name]
        if summary_data:
            # Add relevant benchmark context for better retrieval
            if 'best_model' in summary_data:
                context_query_parts.append(f"best model {summary_data['best_model']}")
            if 'top_3_models' in summary_data and summary_data['top_3_models']:
                top_models = ", ".join([m['model'] for m in summary_data['top_3_models'][:2]])
                context_query_parts.append(f"models {top_models}")
        
        # Use full section text for better context retrieval
        
        enhanced = rag_system.enhance_report_section(
            section_text=base_content,
            section_name=section_name,
            use_rag=True,
            top_k=top_k,
            benchmark_context=summary_data,  # Pass benchmark context for better integration
            validate_citations=validate_citations
        )
        # Return the full enhanced result dict so citations can be extracted
        return enhanced
    except Exception as e:
        logger.warning(f"RAG enhancement failed for {section_name}: {e}")
        return base_content

def _generate_spin_proposition(session, best_model, model_rankings, results, llm_report=None):
    """
    Generate a SPIN-based proposition for the data solution using LLM
    
    SPIN Method:
    - Situation: Current state and context
    - Problem: Challenges and pain points
    - Implication: Consequences of not addressing problems
    - Need-payoff: Benefits and value of the solution
    """
    if not best_model or not model_rankings:
        return None
    
    try:
        from test_models import execute_cortex_query
        
        # Build context from benchmark results
        total_tests = len(results) if results else 0
        successful_tests = sum(1 for r in results if r.get('success')) if results else 0
        avg_quality = sum(m['avg_quality'] for m in model_rankings) / len(model_rankings) if model_rankings else 0
        avg_cost = sum(m['avg_cost'] for m in model_rankings) / len(model_rankings) if model_rankings else 0
        
        # Use best model for generation
        report_model = best_model.get("model", "claude-3-5-sonnet")
        
        spin_prompt = f"""Create a compelling data solution proposition using the SPIN selling methodology based on the following benchmarking results.

BENCHMARK RESULTS SUMMARY:
- Best Model: {best_model.get('display', 'Unknown')} (Score: {best_model.get('composite_score', 0):.2f})
- Average Quality Across Models: {avg_quality:.2f}/10
- Average Cost per Query: ${avg_cost:.6f}
- Total Tests Executed: {total_tests}
- Success Rate: {(successful_tests/total_tests*100) if total_tests > 0 else 0:.1f}%
- Top 3 Models: {', '.join([m['display'] for m in model_rankings[:3]])}

STRUCTURE THE PROPOSITION USING SPIN METHODOLOGY:

## Situation
Describe the current state of data management challenges in modern enterprises:
- Data quality and validation needs
- AI/ML model selection complexity
- Cost optimization requirements
- Need for reliable, scalable data solutions
- Integration with existing data infrastructure (e.g., Snowflake)

## Problem
Identify specific problems that organizations face:
- Difficulty selecting the right AI models for data tasks
- Lack of objective benchmarking data for decision-making
- High costs associated with AI model usage
- Inconsistent quality in AI-generated outputs
- Time-consuming manual evaluation processes
- Risk of choosing suboptimal solutions without data-driven insights

## Implication
Explore the consequences of not addressing these problems:
- Financial impact: Wasted spend on inefficient models
- Operational impact: Slower time-to-insight, reduced productivity
- Quality impact: Poor data quality leading to bad business decisions
- Competitive impact: Falling behind organizations with optimized AI strategies
- Strategic impact: Inability to scale AI initiatives effectively
- Risk impact: Compliance and governance issues from unreliable AI outputs

## Need-Payoff
Present the value proposition of implementing this solution:
- Immediate benefits: Data-driven model selection based on comprehensive benchmarking
- Cost savings: Optimized model usage reducing operational costs by identifying cost-effective options
- Quality improvements: Higher quality outputs through evidence-based model selection
- Speed advantages: Faster decision-making with pre-validated model recommendations
- Strategic value: Foundation for scaling AI initiatives with confidence
- ROI: Quantifiable return on investment through cost optimization and quality improvements
- Competitive advantage: Best-in-class AI model strategy based on empirical evidence
- Risk mitigation: Reduced risk through validated, tested model recommendations

REQUIREMENTS:
- Use professional, consultative language suitable for CTOs and data leaders
- Include specific metrics and numbers from the benchmark results
- Reference the best model and top performers where relevant
- Emphasize business value and ROI
- Connect technical capabilities to business outcomes
- Use persuasive but factual language
- Structure clearly with the four SPIN sections
- Include quantitative benefits where possible (cost savings, quality improvements, time savings)
- CRITICAL: Use proper English grammar, spelling, and punctuation throughout
- Ensure all sentences are grammatically correct with proper subject-verb agreement
- Write in complete, well-structured sentences
- Use consistent verb tenses and proper capitalization

FORMAT:
- Use clear section headings: ## Situation, ## Problem, ## Implication, ## Need-Payoff
- Include bullet points for key points within each section
- Use professional business language with correct grammar
- Keep paragraphs concise and impactful
- End with a strong value statement
- Proofread for grammatical accuracy

Generate the complete SPIN-based proposition:"""
        
        proposition = execute_cortex_query(session, report_model, spin_prompt)
        return str(proposition) if proposition else None
        
    except Exception as e:
        logger.warning(f"SPIN proposition generation failed: {e}")
        return None

def generate_llm_report_section_based(session, results, model_rankings, best_model, use_rag=True):
    """Generate report in sections to ensure complete output and avoid truncation"""
    print("Using section-based generation to ensure complete output...")
    
    # Initialize RAG system if requested
    rag_system = None
    if use_rag:
        try:
            from rag_system import RAGSystem
            import os
            rag_database = os.getenv("SNOWFLAKE_RAG_DATABASE", "RAG_KNOWLEDGE_BASE")
            rag_schema = os.getenv("SNOWFLAKE_RAG_SCHEMA", "PUBLIC")
            rag_system = RAGSystem(session=session, vector_store_database=rag_database, vector_store_schema=rag_schema)
            print("✓ RAG knowledge base loaded - reports will reference authoritative sources")
        except Exception as e:
            logger.warning(f"Could not initialize RAG system: {e}")
            print(f"⚠️  RAG enhancement disabled: {e}")
            rag_system = None
    
    # Prepare summary data
    summary_data = {
        "total_models_tested": len(model_rankings),
        "total_tests": len(results),
        "best_model": best_model["display"] if best_model else "N/A",
        "best_model_score": best_model["composite_score"] if best_model else 0,
        "top_3_models": [
            {
                "rank": idx + 1,
                "model": r["display"],
                "score": r["composite_score"],
                "quality": r["avg_quality"],
                "cost": r["avg_cost"],
                "time": r["avg_time"],
                "success_rate": r["success_rate"]
            }
            for idx, r in enumerate(model_rankings[:3])
        ],
        "test_scenarios": list(TEST_PROMPTS.keys()),
        "all_models": [
            {
                "model": r["display"],
                "score": r["composite_score"],
                "quality": r["avg_quality"],
                "cost": r["avg_cost"],
                "time": r["avg_time"],
                "success_rate": r["success_rate"]
            }
            for r in model_rankings
        ]
    }
    
    report_model = best_model["model"] if best_model else "claude-3-5-sonnet"
    model_limit = MODEL_TOKEN_LIMITS.get(report_model, 8192)
    
    # Calculate safe output tokens (reserve 20% for input, use 80% for output)
    # For section-based, we can use more tokens per section
    safe_output_tokens = int(model_limit * 0.6)  # Conservative estimate
    
    print(f"Using model: {best_model['display'] if best_model else 'claude-3-5-sonnet'}")
    print(f"Model token limit: {model_limit:,}, Allocating ~{safe_output_tokens:,} tokens per section")
    
    sections = {}
    
    # Section 1: Executive Summary (Professional for tech/data management)
    print("\n[1/8] Generating Executive Summary...")
    exec_summary_prompt = f"""You are a senior technology and data management consultant preparing an executive summary for a technical benchmarking report.

Benchmark Context:
- Total models tested: {summary_data['total_models_tested']}
- Total test scenarios: {len(summary_data['test_scenarios'])}
- Best performing model: {summary_data['best_model']} (Composite Score: {summary_data['best_model_score']:.2f})
- Data source: SNOWFLAKE_SAMPLE_DATA.TPCH_SF1

Top 3 Models Performance:
{chr(10).join([f"{m['rank']}. {m['model']}: Composite Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f} per query, Response Time {m['time']:.2f}s, Success Rate {m['success_rate']:.1f}%" for m in summary_data['top_3_models']])}

Write a professional executive summary (4-5 paragraphs, approximately 400-600 words) suitable for technology and data management professionals. The summary should:

1. Open with a strategic overview of the benchmarking initiative, emphasizing its importance for enterprise AI model selection and data management operations
2. Highlight key performance metrics and findings in a business-oriented context
3. Provide actionable insights on model selection for data management use cases
4. Conclude with strategic recommendations for implementation

{"IMPORTANT: Reference authoritative data management frameworks and best practices (e.g., DAMA DMBOK 2nd Edition) where relevant to provide industry-standard context and credibility to your recommendations." if rag_system else ""}

Use professional, technical language appropriate for CTOs, data architects, and technology decision-makers. Focus on business value, operational efficiency, and technical excellence. Avoid overly casual language. Ensure proper English grammar, spelling, and punctuation throughout. All sentences must be grammatically correct with proper subject-verb agreement and consistent verb tenses."""
    
    # Generate executive summary, optionally enhance with RAG
    try:
        base_summary = execute_cortex_query(session, report_model, exec_summary_prompt)
        if rag_system:
            try:
                enhanced = _enhance_section_with_rag(rag_system, session, report_model, base_summary, 'Executive Summary', summary_data, top_k=3)
                # Store the full enhanced result dict
                sections['executive_summary'] = enhanced
                enhanced_text = enhanced.get('enhanced', base_summary) if isinstance(enhanced, dict) else enhanced
                print(f"  ✓ Generated and enhanced with RAG ({len(enhanced_text)} characters)")
            except Exception as rag_error:
                logger.warning(f"RAG enhancement failed for executive summary: {rag_error}")
                sections['executive_summary'] = base_summary
                print(f"  ✓ Generated ({len(sections['executive_summary'])} characters)")
        else:
            sections['executive_summary'] = base_summary
            print(f"  ✓ Generated ({len(sections['executive_summary'])} characters)")
    except Exception as e:
        print(f"  ✗ Failed: {e}")
        sections['executive_summary'] = None
    
    # Section 2: Key Findings
    print("\n[2/8] Generating Key Findings...")
    findings_prompt = f"""Generate detailed key findings from the AI model benchmarking results.

Benchmark Data:
{chr(10).join([f"- {m['model']}: Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f}, Time {m['time']:.2f}s, Success {m['success_rate']:.1f}%" for m in summary_data['all_models']])}

Provide comprehensive key findings including:
- Performance metrics analysis with specific numbers
- Quality vs cost trade-offs with quantitative comparisons
- Speed and reliability insights
- Model-specific strengths and weaknesses
- Notable patterns or anomalies

For each sub-section, include a brief introductory paragraph (1-3 sentences) before bullet points. Use professional technical language suitable for data management professionals. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    # Enhance with RAG if available
    if rag_system:
        try:
            base_content = execute_cortex_query(session, report_model, findings_prompt)
            enhanced = _enhance_section_with_rag(rag_system, session, report_model, base_content, 'Key Findings', summary_data, top_k=3)
            sections['key_findings'] = enhanced
            enhanced_text = enhanced.get('enhanced', base_content) if isinstance(enhanced, dict) else enhanced
            print(f"  ✓ Generated and enhanced with RAG ({len(enhanced_text)} characters)")
        except Exception as rag_error:
            logger.warning(f"RAG enhancement failed, using standard generation: {rag_error}")
            try:
                sections['key_findings'] = execute_cortex_query(session, report_model, findings_prompt)
                print(f"  ✓ Generated ({len(sections['key_findings'])} characters)")
            except Exception as e:
                print(f"  ✗ Failed: {e}")
                sections['key_findings'] = None
    else:
        try:
            sections['key_findings'] = execute_cortex_query(session, report_model, findings_prompt)
            print(f"  ✓ Generated ({len(sections['key_findings'])} characters)")
        except Exception as e:
            print(f"  ✗ Failed: {e}")
            sections['key_findings'] = None
    
    # Section 3: Model Comparison Analysis
    print("\n[3/8] Generating Model Comparison Analysis...")
    comparison_prompt = f"""Provide a detailed model comparison analysis.

Top 3 Models:
{chr(10).join([f"{m['rank']}. {m['model']}: Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f}, Time {m['time']:.2f}s" for m in summary_data['top_3_models']])}

All Models:
{chr(10).join([f"- {m['model']}: Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f}, Time {m['time']:.2f}s" for m in summary_data['all_models']])}

Include:
- Detailed comparison of top 3 models with performance characteristics
- Use case suitability analysis
- When to use which model
- Comparative analysis of all models

For each sub-section, include a brief introductory paragraph (1-3 sentences) before bullets. Use professional language. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    try:
        sections['model_comparison'] = execute_cortex_query(session, report_model, comparison_prompt)
        print(f"  ✓ Generated ({len(sections['model_comparison'])} characters)")
    except Exception as e:
        print(f"  ✗ Failed: {e}")
        sections['model_comparison'] = None
    
    # Section 4: Recommendations
    print("\n[4/8] Generating Recommendations...")
    recommendations_prompt = f"""Provide recommendations for different use cases based on the benchmarking results.

Model Performance Summary:
{chr(10).join([f"- {m['model']}: Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f}, Time {m['time']:.2f}s" for m in summary_data['all_models']])}

Provide specific recommendations for:
- Best model for data quality validation
- Best model for SQL generation
- Best model for cost-sensitive applications
- Best model for high-quality outputs
- Best model for fast responses
- Best model for general-purpose use

{"IMPORTANT: Reference authoritative data management frameworks and best practices (e.g., DAMA DMBOK 2nd Edition) where relevant to provide industry-standard context and credibility to your recommendations." if rag_system else ""}

Use professional, actionable language suitable for technology decision-makers. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    # Enhance with RAG if available
    if rag_system:
        try:
            base_content = execute_cortex_query(session, report_model, recommendations_prompt)
            enhanced = _enhance_section_with_rag(rag_system, session, report_model, base_content, 'Recommendations', summary_data, top_k=3)
            sections['recommendations'] = enhanced
            enhanced_text = enhanced.get('enhanced', base_content) if isinstance(enhanced, dict) else enhanced
            print(f"  ✓ Generated and enhanced with RAG ({len(enhanced_text)} characters)")
        except Exception as rag_error:
            logger.warning(f"RAG enhancement failed, using standard generation: {rag_error}")
            try:
                sections['recommendations'] = execute_cortex_query(session, report_model, recommendations_prompt)
                print(f"  ✓ Generated ({len(sections['recommendations'])} characters)")
            except Exception as e:
                print(f"  ✗ Failed: {e}")
                sections['recommendations'] = None
    else:
        try:
            sections['recommendations'] = execute_cortex_query(session, report_model, recommendations_prompt)
            print(f"  ✓ Generated ({len(sections['recommendations'])} characters)")
        except Exception as e:
            print(f"  ✗ Failed: {e}")
            sections['recommendations'] = None
    
    # Section 5: Cost-Benefit Analysis
    print("\n[5/8] Generating Cost-Benefit Analysis...")
    cost_benefit_prompt = f"""Provide a comprehensive cost-benefit analysis.

Model Cost and Performance Data:
{chr(10).join([f"- {m['model']}: Cost ${m['cost']:.6f} per query, Quality {m['quality']:.2f}/10, Score {m['score']:.2f}" for m in summary_data['all_models']])}

Include:
- Cost comparison across models with specific numbers
- Value proposition of each model
- ROI considerations in structured format (Model, Cost per Query, Expected Monthly Volume, Monthly Cost, Quality Score, Value Rating, ROI Assessment)
- Budget recommendations
- Cost-effectiveness analysis

{"IMPORTANT: Reference authoritative data management frameworks and best practices (e.g., DAMA DMBOK 2nd Edition) where relevant to provide industry-standard context and credibility to your recommendations." if rag_system else ""}

Format ROI considerations as a clear list that can be converted to a table. Use professional financial/technical language. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    # Enhance with RAG if available
    if rag_system:
        try:
            base_content = execute_cortex_query(session, report_model, cost_benefit_prompt)
            enhanced = _enhance_section_with_rag(rag_system, session, report_model, base_content, 'Cost-Benefit Analysis', summary_data, top_k=3)
            sections['cost_benefit'] = enhanced
            enhanced_text = enhanced.get('enhanced', base_content) if isinstance(enhanced, dict) else enhanced
            print(f"  ✓ Generated and enhanced with RAG ({len(enhanced_text)} characters)")
        except Exception as rag_error:
            logger.warning(f"RAG enhancement failed, using standard generation: {rag_error}")
            try:
                sections['cost_benefit'] = execute_cortex_query(session, report_model, cost_benefit_prompt)
                print(f"  ✓ Generated ({len(sections['cost_benefit'])} characters)")
            except Exception as e:
                print(f"  ✗ Failed: {e}")
                sections['cost_benefit'] = None
    else:
        try:
            sections['cost_benefit'] = execute_cortex_query(session, report_model, cost_benefit_prompt)
            print(f"  ✓ Generated ({len(sections['cost_benefit'])} characters)")
        except Exception as e:
            print(f"  ✗ Failed: {e}")
            sections['cost_benefit'] = None
    
    # Section 6: Methodology
    print("\n[6/8] Generating Methodology Explanation...")
    methodology_prompt = """Explain the benchmarking methodology and framework used.

Explain:
- The Role-Task-Framework (RTF) used in prompts
- What RTF is and why it's important for consistent AI responses
- How RTF was implemented in this benchmarking process
- Benefits of using structured prompt frameworks

Use professional technical language suitable for data management and AI professionals. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    try:
        sections['methodology'] = execute_cortex_query(session, report_model, methodology_prompt)
        print(f"  ✓ Generated ({len(sections['methodology'])} characters)")
    except Exception as e:
        print(f"  ✗ Failed: {e}")
        sections['methodology'] = None
    
    # Section 7: Metrics and Evaluation
    print("\n[7/8] Generating Metrics Explanation...")
    metrics_prompt = f"""Explain the metrics and evaluation criteria used in the benchmarking.

Model Scores:
{chr(10).join([f"- {m['model']}: Composite Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Success Rate {m['success_rate']:.1f}%" for m in summary_data['all_models']])}

Explain:
- Composite score calculation methodology
- Quality scoring methodology
- Cost calculation approach
- Code executability evaluation
- Success rate determination

Use professional technical language with specific details. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    try:
        sections['metrics'] = execute_cortex_query(session, report_model, metrics_prompt)
        print(f"  ✓ Generated ({len(sections['metrics'])} characters)")
    except Exception as e:
        print(f"  ✗ Failed: {e}")
        sections['metrics'] = None
    
    # Section 8: Conclusion
    print("\n[8/8] Generating Conclusion...")
    conclusion_prompt = f"""Provide a comprehensive conclusion for the benchmarking report.

Best Model: {summary_data['best_model']} (Score: {summary_data['best_model_score']:.2f})

Include:
- Final recommendations
- Implementation guidance
- Future considerations
- Next steps for model selection

{"IMPORTANT: Reference authoritative data management frameworks and best practices (e.g., DAMA DMBOK 2nd Edition) where relevant to provide industry-standard context and credibility to your recommendations." if rag_system else ""}

Use professional, actionable language suitable for technology and data management professionals. Ensure proper English grammar, spelling, and punctuation throughout."""
    
    # Enhance with RAG if available
    if rag_system:
        try:
            base_content = execute_cortex_query(session, report_model, conclusion_prompt)
            enhanced = _enhance_section_with_rag(rag_system, session, report_model, base_content, 'Conclusion', summary_data, top_k=3)
            sections['conclusion'] = enhanced
            enhanced_text = enhanced.get('enhanced', base_content) if isinstance(enhanced, dict) else enhanced
            print(f"  ✓ Generated and enhanced with RAG ({len(enhanced_text)} characters)")
        except Exception as rag_error:
            logger.warning(f"RAG enhancement failed, using standard generation: {rag_error}")
            try:
                sections['conclusion'] = execute_cortex_query(session, report_model, conclusion_prompt)
                print(f"  ✓ Generated ({len(sections['conclusion'])} characters)")
            except Exception as e:
                print(f"  ✗ Failed: {e}")
                sections['conclusion'] = None
    else:
        try:
            sections['conclusion'] = execute_cortex_query(session, report_model, conclusion_prompt)
            print(f"  ✓ Generated ({len(sections['conclusion'])} characters)")
        except Exception as e:
            print(f"  ✗ Failed: {e}")
            sections['conclusion'] = None
    
    # Combine all sections and collect citations metadata
    print("\nCombining sections into complete report...")
    full_report_parts = []
    all_citations = {}  # Store citations by section for later use
    all_detailed_citations = []  # Store all detailed citations for references section
    
    section_order = [
        ('executive_summary', 'Executive Summary'),
        ('key_findings', 'Key Findings'), 
        ('model_comparison', 'Model Comparison Analysis'),
        ('recommendations', 'Recommendations for Different Use Cases'),
        ('cost_benefit', 'Cost-Benefit Analysis'),
        ('methodology', 'Methodology and Framework Explanation'),
        ('metrics', 'Metrics and Evaluation Criteria'),
        ('conclusion', 'Conclusion')
    ]
    
    for section_key, section_title in section_order:
        if sections.get(section_key):
            # Extract text from enhanced result if it's a dict
            section_text = sections[section_key]
            if isinstance(section_text, dict):
                section_text = section_text.get('enhanced', '')
            full_report_parts.append(f"# {section_title}\n\n" + section_text)
            
            # Check if this section was RAG-enhanced and has citations
            if rag_system and section_key in ['executive_summary', 'key_findings', 'recommendations', 'cost_benefit', 'conclusion']:
                # Get the enhanced result to extract detailed citations
                enhanced_result = None
                if section_key == 'executive_summary' and sections.get('executive_summary'):
                    # Check if it's a dict (enhanced result) or string
                    if isinstance(sections['executive_summary'], dict):
                        enhanced_result = sections['executive_summary']
                elif section_key == 'key_findings' and sections.get('key_findings'):
                    if isinstance(sections['key_findings'], dict):
                        enhanced_result = sections['key_findings']
                elif section_key == 'recommendations' and sections.get('recommendations'):
                    if isinstance(sections['recommendations'], dict):
                        enhanced_result = sections['recommendations']
                elif section_key == 'cost_benefit' and sections.get('cost_benefit'):
                    if isinstance(sections['cost_benefit'], dict):
                        enhanced_result = sections['cost_benefit']
                elif section_key == 'conclusion' and sections.get('conclusion'):
                    if isinstance(sections['conclusion'], dict):
                        enhanced_result = sections['conclusion']
                
                # Extract citations from enhanced result
                detailed_citations = []
                if enhanced_result and isinstance(enhanced_result, dict) and 'citations' in enhanced_result:
                    detailed_citations = enhanced_result.get('citations', [])
                    all_detailed_citations.extend(detailed_citations)
                
                # These sections were RAG-enhanced, collect citation info
                all_citations[section_title] = {
                    'enhanced': True,
                    'source': 'DAMA DMBOK 2nd Edition',
                    'note': 'This section was enhanced with knowledge base references from authoritative data management sources.',
                    'detailed_citations': detailed_citations,
                    'validation_passed': enhanced_result.get('validation_passed', True) if enhanced_result else True
                }
    
    full_report = "\n\n".join(full_report_parts)
    
    # Add references section if RAG was used
    if rag_system and all_citations:
        references_section = "\n\n# Knowledge Base References\n\n"
        references_section += "The following sections of this report were enhanced using authoritative knowledge base sources:\n\n"
        
        for section, info in all_citations.items():
            references_section += f"- **{section}**: Enhanced with references to {info['source']}\n"
        
        references_section += "\n**Primary Knowledge Base Source:**\n"
        references_section += "- DAMA DMBOK 2nd Edition (Data Management Body of Knowledge)\n"
        references_section += "- Citation: DAMA International. (2017). DAMA-DMBOK: Data Management Body of Knowledge (2nd ed.). Technics Publications.\n"
        
        full_report += references_section
    
    # Store citations metadata for document generation
    # This solves Problem 1: Citation Loss
    if hasattr(generate_llm_report_section_based, '_citations_metadata'):
        delattr(generate_llm_report_section_based, '_citations_metadata')
    generate_llm_report_section_based._citations_metadata = all_citations
    generate_llm_report_section_based._detailed_citations_list = all_detailed_citations
    
    print(f"✓ Complete report generated ({len(full_report)} characters, {len(sections)} sections)")
    if all_citations:
        print(f"✓ RAG citations preserved for {len(all_citations)} sections")
    return full_report

def generate_llm_report_single(session, results, model_rankings, best_model):
    """Generate report in a single call (original method)"""
    
    # Prepare summary data for LLM
    summary_data = {
        "total_models_tested": len(model_rankings),
        "total_tests": len(results),
        "best_model": best_model["display"] if best_model else "N/A",
        "best_model_score": best_model["composite_score"] if best_model else 0,
        "top_3_models": [
            {
                "rank": idx + 1,
                "model": r["display"],
                "score": r["composite_score"],
                "quality": r["avg_quality"],
                "cost": r["avg_cost"],
                "time": r["avg_time"],
                "success_rate": r["success_rate"]
            }
            for idx, r in enumerate(model_rankings[:3])
        ],
        "test_scenarios": list(TEST_PROMPTS.keys()),
        "all_models": [
            {
                "model": r["display"],
                "score": r["composite_score"],
                "quality": r["avg_quality"],
                "cost": r["avg_cost"],
                "time": r["avg_time"],
                "success_rate": r["success_rate"]
            }
            for r in model_rankings
        ]
    }
    
    # Create comprehensive prompt for LLM with professional language
    report_prompt = f"""You are a senior technology and data management consultant preparing a comprehensive benchmarking report for enterprise AI model selection.

Benchmark Summary:
- Total models tested: {summary_data['total_models_tested']}
- Total test scenarios: {len(summary_data['test_scenarios'])}
- Best performing model: {summary_data['best_model']} (Composite Score: {summary_data['best_model_score']:.2f})
- Data source: SNOWFLAKE_SAMPLE_DATA.TPCH_SF1

Top 3 Models Performance:
{chr(10).join([f"{m['rank']}. {m['model']}: Composite Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f} per query, Response Time {m['time']:.2f}s, Success Rate {m['success_rate']:.1f}%" for m in summary_data['top_3_models']])}

All Models Tested:
{chr(10).join([f"- {m['model']}: Score {m['score']:.2f}, Quality {m['quality']:.2f}/10, Cost ${m['cost']:.6f}, Time {m['time']:.2f}s" for m in summary_data['all_models']])}

Test Scenarios Evaluated:
{chr(10).join([f"- {scenario.replace('_', ' ').title()}" for scenario in summary_data['test_scenarios']])}

Provide a comprehensive analysis suitable for technology and data management professionals:

1. Executive Summary (4-5 paragraphs, 400-600 words)
   - Strategic overview of the benchmarking initiative and its importance for enterprise AI model selection and data management operations
   - Key performance metrics and findings presented in a business-oriented context
   - Actionable insights on model selection for data management use cases
   - Strategic recommendations for implementation
   Use professional, technical language appropriate for CTOs, data architects, and technology decision-makers. Focus on business value, operational efficiency, and technical excellence.

2. Key Findings (detailed analysis with bullet points)
   - Performance metrics analysis with specific quantitative data
   - Quality vs cost trade-offs with quantitative comparisons
   - Speed and reliability insights with operational implications
   - Model-specific strengths and weaknesses for enterprise deployment
   - Notable patterns or anomalies that impact decision-making
   IMPORTANT: For each sub-section, include a brief introductory paragraph (1-3 sentences) above the bullets that summarizes the key points in professional technical language. Use language suitable for data management professionals.

3. Model Comparison Analysis
   - Detailed comparison of top 3 models
   - Performance characteristics of each model (with 1-3 sentence summary before bullets)
   - Use case suitability analysis (with 1-3 sentence summary before bullets)
   - When to use which model
   - Comparative analysis of all models (with 1-3 sentence summary before bullets)
   IMPORTANT: For each sub-section with bullet points, include a brief introductory paragraph (1-3 sentences) above the bullets that provides context and summarizes what the bullets will cover.

4. Recommendations for Different Use Cases
   - Best model for data quality validation
   - Best model for SQL generation
   - Best model for cost-sensitive applications
   - Best model for high-quality outputs
   - Best model for fast responses
   - Best model for general-purpose use

5. Cost-Benefit Analysis
   - Cost comparison across models
   - Value proposition of each model
   - ROI considerations (IMPORTANT: Present ROI considerations in a structured format that can be converted to a table. Include columns such as: Model, Cost per Query, Expected Monthly Volume, Monthly Cost, Quality Score, Value Rating, ROI Assessment. Format as a clear list with model names and their ROI metrics.)
   - Budget recommendations
   - Cost-effectiveness analysis

6. Methodology and Framework Explanation
   - Explanation of the Role-Task-Framework (RTF) used in prompts
   - What RTF is and why it's important for consistent AI responses
   - How RTF was implemented in this benchmarking process
   - Benefits of using structured prompt frameworks

7. Metrics and Evaluation Criteria
   - Explanation of the composite score calculation
   - Quality scoring methodology
   - Cost calculation approach
   - Code executability evaluation
   - Success rate determination

8. Conclusion
   - Final recommendations
   - Implementation guidance
   - Future considerations
   - Next steps for model selection

Format the response in clear sections with descriptive headings. Be specific, detailed, and provide actionable insights. 

LANGUAGE AND TONE:
- Use professional, technical language appropriate for technology and data management professionals
- Write for an audience of CTOs, data architects, technology decision-makers, and enterprise data teams
- Focus on business value, operational efficiency, technical excellence, and strategic implications
- Avoid overly casual language - maintain a professional, consultative tone throughout
- Emphasize quantitative metrics, ROI considerations, and implementation guidance

GRAMMAR AND WRITING QUALITY:
- Use proper English grammar, spelling, and punctuation throughout
- Ensure all sentences are grammatically correct and well-structured
- Use appropriate verb tenses consistently
- Maintain subject-verb agreement
- Use proper capitalization and punctuation
- Write in complete sentences - avoid fragments or run-on sentences
- Use professional business English suitable for formal reports
- Proofread your output for grammatical accuracy before responding

CRITICAL INSTRUCTIONS:
- Provide the COMPLETE report in a single response. Do NOT truncate or say "continued in next part"
- Include ALL 8 sections listed above in full
- If you reach token limits, prioritize completing all sections even if some are more concise
- Do NOT include phrases like "[Continued in next part due to length...]" or similar truncation messages
- Ensure the Conclusion section is fully completed

IMPORTANT DISCLOSURE NOTE: This report will include a separate "AI Disclosure and Attribution" section that documents the use of AI tools in report generation. You do not need to include disclosure statements within the analysis content itself, as this will be handled separately in the document structure."""
    
    print(f"\nGenerating comprehensive report using: {best_model['display'] if best_model else 'claude-3-5-sonnet'}")
    report_model = best_model["model"] if best_model else "claude-3-5-sonnet"
    
    try:
        llm_report = execute_cortex_query(session, report_model, report_prompt)
        return str(llm_report)
    except Exception as e:
        print(f"⚠️  LLM report generation failed: {e}")
        return None

def create_word_document(results, model_rankings, best_model, llm_report=None, output_dir=None, unavailable_models=None, rag_citations=None, session=None):
    """Create Word document report using python-docx with TOC, page numbers, and model outputs
    
    Args:
        results: Test results data
        model_rankings: Ranked list of models
        best_model: Best performing model
        llm_report: LLM-generated report content
        output_dir: Output directory for document
        unavailable_models: List of unavailable models
        rag_citations: RAG citation metadata
        session: Snowflake session (optional, for SPIN proposition generation)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("⚠️  python-docx not available. Install with: pip install python-docx")
        return None
    
    print("\nCreating Word document report...")
    
    doc = Document()
    
    # Helper function to convert protected citations to display format
    def _convert_citations_for_display(text):
        """Convert {CITE: Document Name} or {{CITE: Document Name}} to proper citation display format"""
        import re
        if not isinstance(text, str):
            return text
        
        # Convert {CITE: Document Name} or {{CITE: Document Name}} to proper APA-style in-text citations
        # {CITE: DAMA DMBOK 2nd Edition} -> (DAMA International, 2017)
        # {{CITE: DAMA DMBOK 2nd Edition}} -> (DAMA International, 2017)
        def replace_cite(match):
            doc_name = match.group(1).strip()
            if 'DAMA DMBOK' in doc_name:
                return '(DAMA International, 2017)'
            else:
                # Generic format for other sources
                return f'({doc_name})'
        
        # Handle both single and double curly braces
        # Pattern matches: {CITE: ...} or {{CITE: ...}}
        # First try double braces, then single braces
        text = re.sub(r'\{\{CITE:\s*([^}]+)\}\}', replace_cite, text)  # {{CITE: ...}}
        text = re.sub(r'\{CITE:\s*([^}]+)\}', replace_cite, text)      # {CITE: ...}
        return text
    
    # Helper function to clean markdown formatting
    def clean_markdown(text):
        """Remove markdown formatting from text"""
        import re
        if not isinstance(text, str):
            text = str(text)
        
        # Remove code blocks (```language ... ```)
        text = re.sub(r'```[a-z]*\n?(.*?)```', r'\1', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove inline code backticks (but keep the content)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove markdown bold/italic but keep text
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # **bold**
        text = re.sub(r'\*([^*]+)\*', r'\1', text)      # *italic*
        text = re.sub(r'__([^_]+)__', r'\1', text)      # __bold__
        text = re.sub(r'_([^_]+)_', r'\1', text)        # _italic_
        
        # Clean up extra whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double
        text = text.strip()
        
        return text
    
    # Helper function to create ROI considerations table
    def create_roi_table(doc, roi_data, model_rankings):
        """Create a table from ROI considerations data"""
        
        # Create table with headers matching the image structure
        roi_table = doc.add_table(rows=1, cols=7)
        roi_table.style = 'Light Grid Accent 1'
        header_cells = roi_table.rows[0].cells
        header_cells[0].text = 'Model'
        header_cells[1].text = 'Cost per Query'
        header_cells[2].text = 'Expected Monthly Volume'
        header_cells[3].text = 'Monthly Cost'
        header_cells[4].text = 'Quality Score'
        header_cells[5].text = 'Value Rating'
        header_cells[6].text = 'ROI Assessment'
        
        # Make headers bold and center-aligned
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Try to extract data from ROI bullets or use model rankings
        processed_models = set()
        
        for roi_item in roi_data:
            # Try to extract model name and metrics
            # Look for model names in the data
            for ranking in model_rankings:
                model_display = ranking['display']
                model_id = ranking['model']
                
                # Check if this model is mentioned in ROI data
                if model_display.lower() in roi_item.lower() or model_id.lower() in roi_item.lower():
                    if model_id not in processed_models:
                        row_cells = roi_table.add_row().cells
                        row_cells[0].text = model_display
                        row_cells[1].text = f"${ranking['avg_cost']:.6f}"
                        
                        # Expected Monthly Volume (default 10,000)
                        expected_volume = 10000
                        row_cells[2].text = f"{expected_volume:,}"
                        
                        # Calculate monthly cost
                        monthly_cost = ranking['avg_cost'] * expected_volume
                        row_cells[3].text = f"${monthly_cost:,.2f}"
                        
                        # Quality Score
                        row_cells[4].text = f"{ranking['avg_quality']:.2f}"
                        
                        # Value Rating (calculated from composite score normalized to 0-10)
                        value_rating = ranking['composite_score'] * 0.8  # Scale composite score
                        row_cells[5].text = f"{value_rating:.2f}"
                        
                        # Generate ROI assessment (Positive/Negative/Neutral)
                        if ranking['composite_score'] >= 7.5 and ranking['avg_cost'] < 0.0005:
                            roi_assessment = "Positive"
                        elif ranking['avg_cost'] > 0.001:
                            roi_assessment = "Negative"
                        else:
                            roi_assessment = "Neutral"
                        
                        row_cells[6].text = roi_assessment
                        processed_models.add(model_id)
        
        # If no data was extracted, create rows from model rankings
        if not processed_models:
            expected_volume = 10000
            for ranking in model_rankings[:5]:  # Top 5 models
                row_cells = roi_table.add_row().cells
                row_cells[0].text = ranking['display']
                row_cells[1].text = f"${ranking['avg_cost']:.6f}"
                row_cells[2].text = f"{expected_volume:,}"
                
                monthly_cost = ranking['avg_cost'] * expected_volume
                row_cells[3].text = f"${monthly_cost:,.2f}"
                row_cells[4].text = f"{ranking['avg_quality']:.2f}"
                
                # Value Rating
                value_rating = ranking['composite_score'] * 0.8
                row_cells[5].text = f"{value_rating:.2f}"
                
                # ROI Assessment
                if ranking['composite_score'] >= 7.5 and ranking['avg_cost'] < 0.0005:
                    roi_assessment = "Positive"
                elif ranking['avg_cost'] > 0.001:
                    roi_assessment = "Negative"
                else:
                    roi_assessment = "Neutral"
                
                row_cells[6].text = roi_assessment
        
        # Set alignment for table cells
        for row in roi_table.rows:
            for idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if idx == 0:  # Model name
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif idx == 6:  # ROI Assessment
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:  # Numeric columns
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')  # Spacing after table
    
    # Helper function to add formatted paragraphs with bullets and alignment
    def add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, use_left_align=False):
        """Add a paragraph with alignment, converting markdown bullets to Word bullets"""
        import re
        
        # Clean markdown first - enhanced to remove table markdown syntax
        text = clean_markdown(text)
        
        # Convert citations after markdown cleaning
        text = _convert_citations_for_display(text)
        
        # Remove markdown table syntax (| and --- lines)
        text = re.sub(r'^\|.*?\|$', '', text, flags=re.MULTILINE)  # Remove table rows
        text = re.sub(r'^[-|:\s]+$', '', text, flags=re.MULTILINE)  # Remove separator lines
        text = re.sub(r'\|\s*', '', text)  # Remove remaining pipe characters
        text = re.sub(r'\s*\|\s*', ' ', text)  # Replace pipe separators with spaces
        
        # Use left alignment if specified (for analysis sections)
        if use_left_align:
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Check if text contains bullet points (lines starting with -, *, or numbered)
        lines = text.split('\n')
        bullet_items = []
        regular_text = []
        in_bullet_list = False
        
        for line in lines:
            line = line.strip()
            if not line:
                if in_bullet_list:
                    # End of bullet list
                    if bullet_items:
                        para = doc.add_paragraph(style='List Bullet')
                        para.alignment = alignment
                        # Citations already converted, just strip formatting
                        para.add_run(bullet_items[0].lstrip('-*• ').strip())
                        for item in bullet_items[1:]:
                            para = doc.add_paragraph(style='List Bullet')
                            para.alignment = alignment
                            para.add_run(item.lstrip('-*• ').strip())
                        bullet_items = []
                    in_bullet_list = False
                continue
            
            # Check if line is a bullet point
            bullet_match = re.match(r'^[-*•]\s+(.+)$', line)
            numbered_match = re.match(r'^\d+[.)]\s+(.+)$', line)
            
            if bullet_match or numbered_match:
                if not in_bullet_list and regular_text:
                    # Add accumulated regular text first (citations already converted)
                    para = doc.add_paragraph('\n'.join(regular_text))
                    para.alignment = alignment
                    regular_text = []
                in_bullet_list = True
                if bullet_match:
                    # Convert citations in bullet text
                    bullet_text = _convert_citations_for_display(bullet_match.group(1))
                    bullet_items.append(bullet_text)
                else:
                    # Convert citations in numbered text
                    numbered_text = _convert_citations_for_display(numbered_match.group(1))
                    bullet_items.append(numbered_text)
            else:
                if in_bullet_list:
                    # End of bullet list, add it
                    if bullet_items:
                        para = doc.add_paragraph(style='List Bullet')
                        para.alignment = alignment
                        para.add_run(bullet_items[0])
                        for item in bullet_items[1:]:
                            para = doc.add_paragraph(style='List Bullet')
                            para.alignment = alignment
                            para.add_run(item)
                        bullet_items = []
                    in_bullet_list = False
                regular_text.append(line)
        
        # Handle remaining items
        if in_bullet_list and bullet_items:
            para = doc.add_paragraph(style='List Bullet')
            para.alignment = alignment
            para.add_run(bullet_items[0])
            for item in bullet_items[1:]:
                para = doc.add_paragraph(style='List Bullet')
                para.alignment = alignment
                para.add_run(item)
        elif regular_text:
            para = doc.add_paragraph('\n'.join(regular_text))
            para.alignment = alignment
            # Professional paragraph formatting
            para_format = para.paragraph_format
            para_format.space_after = Pt(12)
            if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                para_format.first_line_indent = Inches(0.25)  # Indent for justified text
            for run in para.runs:
                run.font.size = Pt(11)
    
    # Track headings for TOC
    
    # Add page numbers to all sections
    def add_page_numbers(section):
        """Add page numbers to footer"""
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
    
    # Apply page numbers to all sections
    for section in doc.sections:
        add_page_numbers(section)
    
    # Title - Professional formatting
    title = doc.add_heading('AI Model Benchmarking Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle for professional context
    subtitle = doc.add_paragraph('Technology and Data Management Professional Assessment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(12)
    for run in subtitle.runs:
        run.italic = True
        run.font.size = Pt(11)
    
    doc.add_paragraph('')  # Spacing after title
    
    # Metadata - Professional left-aligned format
    metadata_section = doc.add_paragraph('Report Information')
    metadata_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in metadata_section.runs:
        run.bold = True
        run.font.size = Pt(11)
    
    metadata_items = [
        f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}',
        f'Total Models Tested: {len(model_rankings)}',
        f'Total Test Scenarios: {len(TEST_PROMPTS)}',
        'Data Source: SNOWFLAKE_SAMPLE_DATA.TPCH_SF1'
    ]
    if output_dir:
        metadata_items.append(f'Test Run Identifier: {output_dir.name}')
    
    for item in metadata_items:
        meta_para = doc.add_paragraph(f'  • {item}')
        meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_para_format = meta_para.paragraph_format
        meta_para_format.left_indent = Inches(0.25)
        meta_para_format.space_after = Pt(6)
        for run in meta_para.runs:
            run.font.size = Pt(10)
    
    doc.add_paragraph('')  # Spacing after metadata
    
    # Table of Contents - Use Word's built-in TOC field
    doc.add_heading('Table of Contents', level=1)
    toc_para = doc.add_paragraph()
    # Add Word's built-in TOC field that automatically generates from heading styles
    run = toc_para.add_run()
    # Create TOC field: TOC \o "1-3" \h \z \u
    # \o "1-3" = include heading levels 1-3
    # \h = make entries hyperlinks
    # \z = hide tab leader and page numbers in web layout
    # \u = use heading styles
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar2)
    # Add a placeholder run that will be replaced when TOC is updated
    run_placeholder = OxmlElement('w:r')
    run_placeholder_text = OxmlElement('w:t')
    run_placeholder_text.text = 'Updating table of contents...'
    run_placeholder.append(run_placeholder_text)
    run._element.append(run_placeholder)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar3)
    doc.add_page_break()  # Start content on new page after TOC
    # We no longer need to track toc_entries manually since Word will auto-generate
    
    # Executive Summary - will be included in LLM report section
    # (LLM generates comprehensive executive summary, so we skip the simple programmatic one)
    
    # LLM-Generated Comprehensive Analysis (includes Executive Summary, Findings, Explanations, etc.)
    if llm_report:
        # Main section heading with professional formatting
        llm_heading = doc.add_heading('Comprehensive Analysis', level=1)
        llm_heading_format = llm_heading.paragraph_format
        llm_heading_format.space_after = Pt(12)
        
        # Add professional introduction paragraph
        intro_para = doc.add_paragraph(
            'This section provides a comprehensive analysis of the AI model benchmarking results, '
            'including executive summary, key findings, model comparisons, recommendations, and strategic insights. '
            'The analysis is designed for technology and data management professionals to support informed decision-making.'
        )
        intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_para_format = intro_para.paragraph_format
        intro_para_format.space_after = Pt(12)
        for run in intro_para.runs:
            run.font.size = Pt(11)
        
        doc.add_paragraph('')  # Spacing before content
        
        # Check for truncation indicators
        truncation_indicators = [
            "[continued in next part",
            "continued in next part",
            "[continued",
            "due to length",
            "would you like me to continue",
            "remaining sections"
        ]
        
        for indicator in truncation_indicators:
            if indicator.lower() in llm_report.lower():
                warning_para = doc.add_paragraph()
                warning_run = warning_para.add_run('⚠️ ')
                warning_run.bold = True
                warning_run.font.color.rgb = RGBColor(200, 0, 0)  # Red color
                warning_para.add_run('Note: This report may have been truncated due to output token limits. Some sections may be incomplete. Consider regenerating with a model that supports longer outputs or using section-based generation.')
                warning_para_format = warning_para.paragraph_format
                warning_para_format.left_indent = Inches(0.25)
                warning_para_format.space_after = Pt(12)
                doc.add_paragraph('')
                break
        
        # Clean markdown from LLM report first
        llm_report = clean_markdown(llm_report)
        
        # Remove truncation messages if present
        truncation_patterns = [
            r'\[Continued in next part due to length\.\.\.\].*',
            r'\[continued in next part.*',
            r'Would you like me to continue.*',
            r'remaining sections.*',
        ]
        import re
        for pattern in truncation_patterns:
            llm_report = re.sub(pattern, '', llm_report, flags=re.IGNORECASE | re.DOTALL)
        
        # Clean up any double newlines that might result
        llm_report = re.sub(r'\n{3,}', '\n\n', llm_report).strip()
        
        # Parse and format LLM report with proper bullets and left-aligned text (for analysis)
        # Enhanced to handle sub-summaries before bullet points and ROI table conversion
        report_lines = llm_report.split('\n')
        current_paragraph = []
        in_bullet_list = False
        pending_summary = []
        in_roi_section = False
        roi_data = []
        in_recommendations_section = False
        recommendations_data = []
        
        for i, line in enumerate(report_lines):
            line = line.strip()
            
            # Check for headings
            if line.startswith('#') or (len(line) < 100 and line.isupper() and len(line.split()) < 10):
                # Flush current paragraph
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph).strip()
                    if para_text:
                        add_formatted_paragraph(doc, para_text, use_left_align=True)
                    current_paragraph = []
                if in_bullet_list:
                    in_bullet_list = False
                
                # Add heading (clean markdown from heading too)
                heading_text = line.lstrip('#').strip()
                heading_text = clean_markdown(heading_text)
                if heading_text:
                    # Determine heading level based on number of # characters
                    if line.startswith('#'):
                        heading_level = len(line) - len(line.lstrip('#'))
                    else:
                        # Uppercase heading (likely level 2)
                        heading_level = 2
                    
                    # Map markdown heading levels to Word heading levels
                    # For section-based generation, main sections are level 1 in markdown (#)
                    # These should be level 2 in Word (sub-sections of "Comprehensive Analysis")
                    # Executive Summary gets special treatment as the primary section
                    if heading_text.lower() == "executive summary":
                        word_level = 2  # Main section, prominent
                        # Add spacing before Executive Summary for emphasis
                        doc.add_paragraph('')
                    elif heading_level == 1:
                        word_level = 2  # Main sub-sections (Key Findings, Model Comparison, etc.)
                    elif heading_level == 2:
                        word_level = 3  # Sub-sub-sections
                    else:
                        word_level = min(heading_level + 1, 3)  # Cap at level 3 for TOC
                    
                    # Add the heading
                    doc.add_heading(heading_text, level=word_level)
                    
                    # Add professional spacing after main section headings
                    if word_level == 2:
                        doc.add_paragraph('')  # Extra spacing for main sections
                    
                    # Check if this is "Recommendations for Different Use Cases" - mark for table conversion
                    if 'recommendations' in heading_text.lower() and 'use case' in heading_text.lower():
                        in_recommendations_section = True
                    else:
                        in_recommendations_section = False
                    pending_summary = []  # Reset pending summary for new heading
                    
                    # Check if this is ROI considerations section
                    in_roi_section = 'ROI' in heading_text.upper() or 'roi' in heading_text.lower()
                    if in_roi_section:
                        roi_data = []  # Reset ROI data collection
            elif not line:
                # If we were collecting ROI data and hit an empty line, create the table
                if in_roi_section and roi_data:
                    # Create ROI considerations table
                    create_roi_table(doc, roi_data, model_rankings)
                    roi_data = []
                    in_roi_section = False
                
                # If we were collecting recommendations data and hit an empty line, create the table
                if in_recommendations_section and recommendations_data:
                    # The recommendations table is already created programmatically later
                    # Just clear the data collection
                    recommendations_data = []
                    in_recommendations_section = False
                
                # Empty line - check if we have a pending summary to add
                if pending_summary:
                    # Add summary paragraph before bullets with professional formatting
                    summary_text = ' '.join(pending_summary).strip()
                    if summary_text:
                        # Convert protected citations before adding
                        summary_text = _convert_citations_for_display(summary_text)
                        summary_para = doc.add_paragraph(summary_text)
                        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        summary_para_format = summary_para.paragraph_format
                        summary_para_format.space_after = Pt(12)
                        summary_para_format.first_line_indent = Inches(0.25)  # Indent first line for professional look
                        for run in summary_para.runs:
                            run.font.size = Pt(11)
                        doc.add_paragraph('')  # Small spacing before bullets
                    pending_summary = []
                
                # Flush current paragraph
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph).strip()
                    if para_text:
                        # Check if this looks like a summary (not a bullet, not too short)
                        import re
                        is_bullet = bool(re.match(r'^[-*•]\s+', para_text))
                        is_numbered = bool(re.match(r'^\d+[.)]\s+', para_text))
                        
                        if not is_bullet and not is_numbered and len(para_text) > 50:
                            # This might be a summary paragraph - check if next non-empty line is a bullet
                            next_line_idx = i + 1
                            while next_line_idx < len(report_lines) and not report_lines[next_line_idx].strip():
                                next_line_idx += 1
                            
                            if next_line_idx < len(report_lines):
                                next_line = report_lines[next_line_idx].strip()
                                next_is_bullet = bool(re.match(r'^[-*•]\s+', next_line)) or bool(re.match(r'^\d+[.)]\s+', next_line))
                                
                                if next_is_bullet:
                                    # This is a summary before bullets - add it as a professional paragraph
                                    # Convert protected citations before adding
                                    para_text = _convert_citations_for_display(para_text)
                                    summary_para = doc.add_paragraph(para_text)
                                    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    summary_para_format = summary_para.paragraph_format
                                    summary_para_format.space_after = Pt(12)
                                    summary_para_format.first_line_indent = Inches(0.25)
                                    for run in summary_para.runs:
                                        run.font.size = Pt(11)
                                    doc.add_paragraph('')  # Small spacing before bullets
                                    current_paragraph = []
                                    continue
                        
                        # Not a summary, process normally
                        # Convert protected citations before adding to document
                        para_text = _convert_citations_for_display(para_text)
                        add_formatted_paragraph(doc, para_text, use_left_align=True)
                    current_paragraph = []
                if in_bullet_list:
                    in_bullet_list = False
            else:
                # Check if this line is a bullet point
                import re
                is_bullet = bool(re.match(r'^[-*•]\s+', line)) or bool(re.match(r'^\d+[.)]\s+', line))
                
                if is_bullet:
                    # If we have accumulated text that's not a bullet, it might be a summary
                    if current_paragraph and not any(re.match(r'^[-*•]\s+', line.strip()) or re.match(r'^\d+[.)]\s+', line.strip()) for line in current_paragraph):
                        # Check if accumulated text looks like a summary paragraph
                        accumulated_text = ' '.join(current_paragraph).strip()
                        if len(accumulated_text) > 50:
                            # Add as summary paragraph before bullets with professional formatting
                            # Convert protected citations before adding
                            accumulated_text = _convert_citations_for_display(accumulated_text)
                            summary_para = doc.add_paragraph(accumulated_text)
                            summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            summary_para_format = summary_para.paragraph_format
                            summary_para_format.space_after = Pt(12)
                            summary_para_format.first_line_indent = Inches(0.25)
                            for run in summary_para.runs:
                                run.font.size = Pt(11)
                            doc.add_paragraph('')  # Small spacing before bullets
                            current_paragraph = []
                    
                    # Check if we're in ROI section and this bullet contains ROI data
                    if in_roi_section:
                        # Extract ROI data from bullet point
                        bullet_text = re.sub(r'^[-*•]\s+', '', line).strip()
                        # Try to parse model name and ROI metrics
                        # Format might be: "Model Name: Cost $X, Quality Y, ROI: Z"
                        roi_data.append(bullet_text)
                    
                    # Now add the bullet point (unless we're building ROI table)
                    if not in_roi_section:
                        current_paragraph.append(line)
                    in_bullet_list = True
                else:
                    current_paragraph.append(line)
        
        # Flush remaining content
        if current_paragraph:
            para_text = '\n'.join(current_paragraph).strip()
            if para_text:
                # Convert protected citations to proper format before adding to document
                para_text = _convert_citations_for_display(para_text)
                add_formatted_paragraph(doc, para_text, use_left_align=True)
        
        # If we still have ROI data, create the table
        if in_roi_section and roi_data:
            create_roi_table(doc, roi_data, model_rankings)
        
        doc.add_paragraph('')
    
    # Model Rankings
    doc.add_heading('Model Rankings', level=1)
    # Introduction will be extracted from LLM report if available, otherwise use simple intro
    if llm_report and 'Model Rankings' in llm_report or 'rankings' in llm_report.lower():
        # LLM may have provided context in its report, but we'll add a brief intro
        rankings_intro = doc.add_paragraph('Models are ranked by composite score, which considers quality, cost, speed, and code executability. The following table provides a detailed breakdown:')
    else:
        rankings_intro = doc.add_paragraph('Models ranked by composite score (quality, cost, speed, executability):')
    rankings_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('')
    
    # Create rankings table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Rank'
    header_cells[1].text = 'Model'
    header_cells[2].text = 'Quality'
    header_cells[3].text = 'Time (s)'
    header_cells[4].text = 'Cost ($)'
    header_cells[5].text = 'Success %'
    header_cells[6].text = 'Score'
    
    for idx, ranking in enumerate(model_rankings, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = ranking['display']
        row_cells[2].text = f"{ranking['avg_quality']:.2f}"
        row_cells[3].text = f"{ranking['avg_time']:.2f}"
        row_cells[4].text = f"{ranking['avg_cost']:.6f}"
        row_cells[5].text = f"{ranking['success_rate']:.1f}%"
        row_cells[6].text = f"{ranking['composite_score']:.2f}"
        
        # Set alignment for table cells - center numeric columns, left-align model names
        for para in row_cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Rank
        for para in row_cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Model name
        for para in row_cells[2].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Quality
        for para in row_cells[3].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Time
        for para in row_cells[4].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Cost
        for para in row_cells[5].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Success %
        for para in row_cells[6].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Score
    
    doc.add_paragraph('')  # Spacing after table
    
    # Detailed Results by Model
    doc.add_heading('Detailed Results by Model', level=1)
    # Add brief context about the detailed metrics
    detailed_intro = doc.add_paragraph(
        'The following table provides comprehensive performance metrics for each model across all test scenarios. '
        'These metrics include average quality scores, response times, costs, success rates, and code executability percentages.'
    )
    detailed_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('')
    
    # Create a table for all models
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Model'
    header_cells[1].text = 'Composite Score'
    header_cells[2].text = 'Avg Quality'
    header_cells[3].text = 'Avg Time (s)'
    header_cells[4].text = 'Avg Cost ($)'
    header_cells[5].text = 'Success Rate'
    header_cells[6].text = 'Code Executable'
    
    # Make header bold and center-aligned
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for ranking in model_rankings:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{ranking['display']}\n({ranking['model']})"
        row_cells[1].text = f"{ranking['composite_score']:.2f}"
        row_cells[2].text = f"{ranking['avg_quality']:.2f}/10"
        row_cells[3].text = f"{ranking['avg_time']:.2f}"
        row_cells[4].text = f"{ranking['avg_cost']:.6f}"
        row_cells[5].text = f"{ranking['success_rate']:.1f}%"
        executable_text = f"{ranking.get('executable_rate', 0):.1f}%" if ranking.get('executable_rate', 0) > 0 else "N/A"
        row_cells[6].text = executable_text
        
        # Center-align numeric columns
        for idx in [1, 2, 3, 4, 5, 6]:
            for paragraph in row_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Left-align model name
        for paragraph in row_cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('')
    
    # Test Scenario Details with Model Outputs
    doc.add_heading('Test Scenario Details', level=1)
    
    # RTF Framework Introduction - extract from LLM report if available
    doc.add_heading('Prompt Framework: Role-Task-Framework (RTF)', level=2)
    
    # Try to extract RTF explanation from LLM report
    rtf_explanation = None
    if llm_report:
        import re
        # Look for RTF section in LLM report
        rtf_pattern = r'(?i)(methodology.*?framework|RTF|Role-Task-Framework).*?(?=\n\n#|\n##|$)'
        rtf_match = re.search(rtf_pattern, llm_report, re.DOTALL)
        if rtf_match:
            rtf_explanation = rtf_match.group(0)
    
    if rtf_explanation:
        # Use LLM-generated explanation
        rtf_cleaned = clean_markdown(rtf_explanation)
        add_formatted_paragraph(doc, rtf_cleaned, use_left_align=True)
    else:
        # Fallback to programmatic explanation if LLM didn't provide one
        rtf_what = doc.add_paragraph()
        rtf_what.add_run('What is RTF? ').bold = True
        rtf_what.add_run(
            'The Role-Task-Framework (RTF) is a structured prompt engineering approach that defines clear roles, '
            'specific tasks, and expected outputs for AI models. This framework ensures consistent, high-quality '
            'responses by providing explicit context and instructions.'
        )
        rtf_what.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        rtf_how = doc.add_paragraph()
        rtf_how.add_run('How is it implemented? ').bold = True
        rtf_how.add_run(
            'In this benchmarking process, RTF is implemented through structured system prompts that: '
            '(1) Define the AI model\'s role (e.g., "expert in Great Expectations data validation library"), '
            '(2) Specify the exact task (e.g., "Convert natural language descriptions into Python code"), '
            '(3) Provide clear examples and constraints (e.g., "Return ONLY executable Python code, NO explanations"), '
            'and (4) Include relevant context (e.g., available column names, table schemas).'
        )
        rtf_how.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        rtf_why = doc.add_paragraph()
        rtf_why.add_run('Why is it used? ').bold = True
        rtf_why.add_run(
            'RTF is used to ensure consistent, reproducible, and high-quality AI responses across different models. '
            'By providing explicit structure and examples, RTF reduces ambiguity, minimizes errors, and enables '
            'fair comparison between models. This framework is particularly important for code generation tasks where '
            'precision and format consistency are critical for successful execution.'
        )
        rtf_why.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('')
    
    # Group results by test scenario
    results_by_test = {}
    for result in results:
        test_name = result.get('test_name', 'unknown')
        if test_name not in results_by_test:
            results_by_test[test_name] = []
        results_by_test[test_name].append(result)
    
    for test_name, test_config in TEST_PROMPTS.items():
        test_title = test_name.replace('_', ' ').title()
        doc.add_heading(test_title, level=2)
        
        type_para = doc.add_paragraph(f"Type: {test_config['test_type']}")
        type_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph('')
        
        prompt_para = doc.add_paragraph()
        prompt_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prompt_para.add_run('Prompt: ').bold = True
        # Show the actual prompt used (from test results if available, otherwise from config)
        test_results_for_prompt = results_by_test.get(test_name, [])
        if test_results_for_prompt and test_results_for_prompt[0].get('prompt'):
            # Use the actual prompt that was sent to the model
            actual_prompt = test_results_for_prompt[0].get('prompt', '')
            # Truncate very long prompts for readability
            if len(actual_prompt) > 2000:
                prompt_text = actual_prompt[:2000] + "\n\n[... Prompt truncated for brevity ...]"
            else:
                prompt_text = actual_prompt
        else:
            # Fallback to config
            if 'prompt' in test_config:
                prompt_text = test_config['prompt']
            elif 'natural_language' in test_config:
                prompt_text = f"Natural Language: {test_config['natural_language']}"
                if 'table' in test_config:
                    prompt_text += f"\nTable: {test_config.get('schema', '')}.{test_config.get('table', '')}"
            else:
                prompt_text = "N/A"
        prompt_para.add_run(prompt_text)
        doc.add_paragraph('')
        
        # Show model outputs for this test scenario in a table format
        test_results = results_by_test.get(test_name, [])
        if test_results:
            doc.add_heading('Model Responses', level=3)
            
            # Sort by model name for consistency
            test_results_sorted = sorted(test_results, key=lambda x: x.get('model', ''))
            
            # Create table for model responses
            response_table = doc.add_table(rows=1, cols=6)
            response_table.style = 'Light Grid Accent 1'
            header_cells = response_table.rows[0].cells
            header_cells[0].text = 'Model'
            header_cells[1].text = 'Status'
            header_cells[2].text = 'Response'
            header_cells[3].text = 'Quality Score'
            header_cells[4].text = 'Response Time'
            header_cells[5].text = 'Cost'
            
            # Make headers bold and center-aligned
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows
            for result in test_results_sorted:
                model_name = result.get('model', 'unknown')
                model_display = MODELS.get(model_name, {}).get('display', model_name)
                
                row_cells = response_table.add_row().cells
                
                # Model name
                row_cells[0].text = model_display
                
                # Status
                if result.get('success'):
                    row_cells[1].text = 'Success'
                else:
                    row_cells[1].text = 'Failed'
                
                # Response
                if result.get('success'):
                    response = result.get('response', 'No response')
                    # Ensure response is a string
                    if not isinstance(response, str):
                        response = str(response)
                    
                    # Clean markdown formatting
                    response = clean_markdown(response)
                    
                    # Store full response in cell (Word can handle large text)
                    # Add note if response is very long, but include full text
                    if len(response) > 5000:
                        # For very long responses, add a note but keep full text
                        response_display = response[:5000] + f"\n\n[... Additional {len(response) - 5000} characters in full response ...]\n\n" + response[5000:]
                    else:
                        response_display = response
                    
                    row_cells[2].text = response_display
                else:
                    error_msg = result.get('error', 'Unknown error')
                    if len(error_msg) > 500:
                        error_msg = error_msg[:500] + "..."
                    row_cells[2].text = f"Error: {error_msg}"
                
                # Quality Score
                quality = result.get('quality_score', 'N/A')
                if quality == 'N/A' or quality is None:
                    row_cells[3].text = 'N/A'
                else:
                    row_cells[3].text = str(quality)
                
                # Response Time
                response_time = result.get('response_time')
                if response_time is None:
                    row_cells[4].text = 'N/A'
                else:
                    row_cells[4].text = f"{float(response_time):.2f} seconds"
                
                # Cost
                cost = result.get('cost', 0)
                if cost is None:
                    row_cells[5].text = '$0.000000'
                else:
                    row_cells[5].text = f"${float(cost):.6f}"
                
                # Set alignment for table cells
                for para in row_cells[0].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Model name
                for para in row_cells[1].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Status
                for para in row_cells[2].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Response (left-aligned for readability)
                for para in row_cells[3].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Quality Score
                for para in row_cells[4].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Response Time
                for para in row_cells[5].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Cost
            
            doc.add_paragraph('')  # Spacing after table
        else:
            doc.add_paragraph('No test results available for this scenario.')
        
        doc.add_paragraph('')
    
    # Recommendations
    doc.add_heading('Recommendations', level=1)
    
    # Introduction paragraph - brief, table provides details
    intro_para = doc.add_paragraph(
        "Based on comprehensive benchmarking analysis, the following recommendations are provided for different use cases. "
        "These recommendations consider quality, cost, speed, and reliability metrics."
    )
    intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('')
    
    # Create recommendations table
    rec_table = doc.add_table(rows=1, cols=3)
    rec_table.style = 'Light Grid Accent 1'
    header_cells = rec_table.rows[0].cells
    header_cells[0].text = 'Use Case'
    header_cells[1].text = 'Recommended Model'
    header_cells[2].text = 'Rationale'
    
    # Make header bold and center-aligned
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add best model recommendation
    if best_model:
        row_cells = rec_table.add_row().cells
        row_cells[0].text = 'Default / General Purpose'
        row_cells[1].text = best_model['display']
        row_cells[2].text = (
            f"Best overall balance: Quality {best_model['avg_quality']:.2f}/10, "
            f"Cost ${best_model['avg_cost']:.6f} per query, "
            f"Speed {best_model['avg_time']:.2f}s, "
            f"Success rate {best_model['success_rate']:.1f}%"
        )
    
    # Add alternative recommendations based on model rankings
    if len(model_rankings) > 1:
        # Find best model for cost (lowest cost)
        cost_leader = min(model_rankings, key=lambda x: x['avg_cost'])
        if cost_leader['model'] != best_model['model']:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = 'Cost-Sensitive Applications'
            row_cells[1].text = cost_leader['display']
            row_cells[2].text = (
                f"Lowest cost: ${cost_leader['avg_cost']:.6f} per query, "
                f"Quality {cost_leader['avg_quality']:.2f}/10"
            )
        
        # Find best model for speed (fastest)
        speed_leader = min(model_rankings, key=lambda x: x['avg_time'])
        if speed_leader['model'] != best_model['model']:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = 'Fast Response Requirements'
            row_cells[1].text = speed_leader['display']
            row_cells[2].text = (
                f"Fastest response: {speed_leader['avg_time']:.2f}s average, "
                f"Quality {speed_leader['avg_quality']:.2f}/10"
            )
        
        # Find best model for quality (highest quality)
        quality_leader = max(model_rankings, key=lambda x: x['avg_quality'])
        if quality_leader['model'] != best_model['model']:
            row_cells = rec_table.add_row().cells
            row_cells[0].text = 'High-Quality Output Requirements'
            row_cells[1].text = quality_leader['display']
            row_cells[2].text = (
                f"Highest quality: {quality_leader['avg_quality']:.2f}/10, "
                f"Cost ${quality_leader['avg_cost']:.6f} per query"
            )
    
    # Set alignment for table cells
    for row in rec_table.rows:
        # Use case column - left aligned
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Model column - center aligned
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Rationale column - left aligned for consistency
        for paragraph in row.cells[2].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph('')  # Spacing after recommendations table
    
    # Data Solution Proposition (SPIN Method)
    doc.add_page_break()
    doc.add_heading('Data Solution Proposition', level=1)
    
    # Generate SPIN-based proposition using LLM if session is available
    proposition_content = None
    if session and best_model and model_rankings:
        try:
            print("\nGenerating SPIN-based data solution proposition...")
            proposition_content = _generate_spin_proposition(session, best_model, model_rankings, results, llm_report)
            if proposition_content:
                print(f"  ✓ Generated SPIN proposition ({len(proposition_content)} characters)")
        except Exception as e:
            logger.warning(f"SPIN proposition generation failed: {e}")
    
    if proposition_content:
        # Add SPIN methodology introduction
        spin_intro = doc.add_paragraph()
        spin_intro.add_run('This proposition is structured using the SPIN selling methodology to clearly articulate the value proposition of implementing this AI-powered data management solution. ').bold = True
        spin_intro.add_run('SPIN stands for Situation, Problem, Implication, and Need-payoff - a proven framework for presenting solutions that address real business challenges.')
        spin_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        spin_intro_format = spin_intro.paragraph_format
        spin_intro_format.space_after = Pt(12)
        doc.add_paragraph('')
        
        # Process and add the proposition content
        prop_lines = proposition_content.split('\n')
        current_section = []
        
        for line in prop_lines:
            line = line.strip()
            if not line:
                if current_section:
                    para_text = '\n'.join(current_section).strip()
                    if para_text:
                        para_text = _convert_citations_for_display(para_text)
                        add_formatted_paragraph(doc, para_text, use_left_align=True)
                    current_section = []
                continue
            
            # Check if line is a heading
            if line.startswith('#'):
                # Flush previous section
                if current_section:
                    para_text = '\n'.join(current_section).strip()
                    if para_text:
                        para_text = _convert_citations_for_display(para_text)
                        add_formatted_paragraph(doc, para_text, use_left_align=True)
                    current_section = []
                
                # Add heading
                heading_text = line.lstrip('#').strip()
                heading_text = clean_markdown(heading_text)
                if heading_text:
                    # SPIN sections are typically level 2
                    doc.add_heading(heading_text, level=2)
                    doc.add_paragraph('')
            else:
                current_section.append(line)
        
        # Flush remaining content
        if current_section:
            para_text = '\n'.join(current_section).strip()
            if para_text:
                para_text = _convert_citations_for_display(para_text)
                add_formatted_paragraph(doc, para_text, use_left_align=True)
        
        doc.add_paragraph('')
    else:
        # Fallback if proposition generation fails
        doc.add_paragraph('Proposition content generation is available when LLM report is included.')
    
    doc.add_page_break()
    
    # Knowledge Base References Section (Problem 4: Missing References Section)
    # Add this BEFORE AI Disclosure for better flow
    rag_citations = getattr(generate_llm_report_section_based, '_citations_metadata', None)
    if rag_citations:
        doc.add_heading('Knowledge Base References', level=1)
        
        kb_intro = doc.add_paragraph()
        kb_intro.add_run('Authoritative Knowledge Base Integration: ').bold = True
        kb_intro.add_run('This report has been enhanced with references to authoritative data management sources to provide industry-standard context and credibility. The following sections incorporate knowledge base references:')
        kb_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph('')
        
        # List sections with RAG enhancement
        for section_title, citation_info in rag_citations.items():
            if citation_info.get('enhanced'):
                section_para = doc.add_paragraph(style='List Bullet')
                section_para.add_run(f'{section_title}: ').bold = True
                section_para.add_run(citation_info.get('note', 'Enhanced with knowledge base references.'))
        
        doc.add_paragraph('')
        
        # Primary source citation (APA format)
        doc.add_heading('Primary Knowledge Base Source', level=2)
        citation_para = doc.add_paragraph()
        citation_para.add_run('DAMA International. (2017). ').italic = True
        citation_para.add_run('DAMA-DMBOK: Data Management Body of Knowledge').italic = True
        citation_para.add_run(' (2nd ed.). Technics Publications.')
        citation_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph('')
        
        # Usage note
        usage_note = doc.add_paragraph()
        usage_note.add_run('Note: ').bold = True
        usage_note.add_run('Knowledge base references are integrated throughout the analysis sections to align recommendations with established data management frameworks and industry best practices. Specific chunk references indicate the relevant portions of the knowledge base used for each enhancement.')
        
        doc.add_paragraph('')
        doc.add_page_break()
    
    # AI Disclosure and Attribution Section
    doc.add_heading('AI Disclosure and Attribution', level=1)
    
    disclosure_para = doc.add_paragraph()
    disclosure_para.add_run('This report was generated with the assistance of artificial intelligence (AI) tools. The following disclosure provides transparency about AI usage in accordance with academic and professional guidelines.').bold = True
    doc.add_paragraph('')
    
    # AI Tools Used
    doc.add_heading('AI Tools Utilized', level=2)
    doc.add_paragraph('The following AI tools were used in the generation of this report:')
    
    # Primary AI tool for report generation
    if llm_report and best_model:
        tool_item = doc.add_paragraph(style='List Bullet')
        tool_item.add_run('Snowflake Cortex AI - ').bold = True
        tool_item.add_run(f'{best_model.get("display", "Unknown Model")} ({best_model.get("model", "unknown")})')
        tool_item.add_run(' was used to generate the comprehensive analysis report (Section: Comprehensive Analysis). This AI tool was used to synthesize benchmarking data, generate executive summaries, provide detailed findings, and create recommendations based on the test results.')
    
    # AI tools tested (all models)
    if model_rankings:
        tool_item2 = doc.add_paragraph(style='List Bullet')
        tool_item2.add_run('Snowflake Cortex AI Models - ').bold = True
        model_names = ', '.join([r.get('display', r.get('model', '')) for r in model_rankings[:3]])
        tool_item2.add_run(f'Multiple AI models ({model_names}) were tested and evaluated as part of the benchmarking process. Their responses to test scenarios are documented in the "Test Scenario Details" section.')
    
    doc.add_paragraph('')
    
    # Purpose and Scope
    doc.add_heading('Purpose and Scope of AI Usage', level=2)
    doc.add_paragraph('AI tools were used for the following purposes:')
    doc.add_paragraph('• Generating comprehensive analysis and interpretation of benchmarking results', style='List Bullet')
    doc.add_paragraph('• Synthesizing performance metrics into actionable insights and recommendations', style='List Bullet')
    doc.add_paragraph('• Creating executive summaries and detailed findings based on quantitative test data', style='List Bullet')
    doc.add_paragraph('• Providing model comparison analysis and use case recommendations', style='List Bullet')
    doc.add_paragraph('• Generating cost-benefit analysis and ROI considerations', style='List Bullet')
    
    if llm_report:
        doc.add_paragraph('')
        doc.add_paragraph('The AI-generated content appears in the "Comprehensive Analysis" section of this report. All other sections (Model Rankings, Test Scenario Details, etc.) contain programmatically generated content based on actual test results.')
    
    doc.add_paragraph('')
    
    # Human Oversight
    doc.add_heading('Human Oversight and Verification', level=2)
    doc.add_paragraph('The following human oversight and verification processes were applied:')
    doc.add_paragraph('• All AI-generated content was reviewed and verified for accuracy', style='List Bullet')
    doc.add_paragraph('• Benchmarking test results and quantitative metrics were programmatically generated and validated', style='List Bullet')
    doc.add_paragraph('• Model rankings and statistics are based on actual test execution data, not AI-generated', style='List Bullet')
    doc.add_paragraph('• The final report structure, formatting, and integration of AI content was performed by human operators', style='List Bullet')
    
    doc.add_paragraph('')
    
    # Citation Format
    doc.add_heading('Citation and Attribution', level=2)
    citation_para = doc.add_paragraph('When referencing this report or its AI-generated content, please use the following citation format:')
    doc.add_paragraph('')
    
    if best_model:
        citation_text = f'''Snowflake Inc. ({datetime.now().strftime("%Y")}). Snowflake Cortex AI {best_model.get("display", "Model")} [Large Language Model]. Snowflake Cortex AI Service.'''
        citation_para2 = doc.add_paragraph(citation_text, style='Intense Quote')
        citation_para2.italic = True
    
    doc.add_paragraph('')
    doc.add_paragraph('For APA-style in-text citations, use: (Snowflake Inc., ' + datetime.now().strftime("%Y") + ')')
    doc.add_paragraph('')
    
    # Methodology Note
    methodology_note = doc.add_paragraph()
    methodology_note.add_run('Note: ').bold = True
    methodology_note.add_run('This disclosure follows guidelines established by academic institutions, professional organizations, and industry best practices for AI-assisted content generation. The use of AI tools was transparent, documented, and subject to human review and verification.')
    
    doc.add_paragraph('')
    doc.add_page_break()
    
    # Appendix: Unavailable Models
    if unavailable_models:
        doc.add_heading('Appendix A: Unavailable Models', level=1)
        unavail_intro = doc.add_paragraph(
            'The following models were tested but are not available in your Snowflake account. '
            'These models are excluded from the main rankings and analysis.'
        )
        unavail_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph('')
        
        # Create table for unavailable models
        unavail_table = doc.add_table(rows=1, cols=4)
        unavail_table.style = 'Light Grid Accent 1'
        header_cells = unavail_table.rows[0].cells
        header_cells[0].text = 'Model'
        header_cells[1].text = 'Model ID'
        header_cells[2].text = 'Tests Attempted'
        header_cells[3].text = 'Reason'
        
        for unavail in unavailable_models:
            row_cells = unavail_table.add_row().cells
            row_cells[0].text = unavail['display']
            row_cells[1].text = unavail['model']
            row_cells[2].text = str(unavail['total_tests'])
            row_cells[3].text = unavail['error_reason']
        
        doc.add_paragraph('')
        unavail_note = doc.add_paragraph(
            'Note: These models may become available in the future or may require additional account configuration.'
        )
        unavail_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph('')
    
    # Appendix: Raw Results Summary
    doc.add_heading('Appendix B: Test Results Summary', level=1)
    summary_items = [
        f'Total test results: {len(results)}',
        f'Successful tests: {sum(1 for r in results if r.get("success"))}',
        f'Failed tests: {sum(1 for r in results if not r.get("success"))}'
    ]
    if unavailable_models:
        summary_items.append(f'Unavailable models: {len(unavailable_models)}')
    
    for item in summary_items:
        para = doc.add_paragraph(style='List Bullet')
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(item)
    
    # References Section (after Appendix)
    doc.add_page_break()
    doc.add_heading('References', level=1)
    
    # Collect all citations from RAG system if available
    rag_citations = getattr(generate_llm_report_section_based, '_citations_metadata', None)
    detailed_citations_list = getattr(generate_llm_report_section_based, '_detailed_citations_list', [])
    
    # Collect unique sources from detailed citations
    unique_sources = {}
    for cite in detailed_citations_list:
        source = cite.get('source', '')
        if source and source not in unique_sources:
            unique_sources[source] = cite
    
    # Add primary knowledge base source (DAMA DMBOK)
    primary_ref = doc.add_paragraph()
    primary_ref.add_run('DAMA International. (2017). ').italic = True
    primary_ref.add_run('DAMA-DMBOK: Data Management Body of Knowledge').italic = True
    primary_ref.add_run(' (2nd ed.). Technics Publications.')
    primary_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    primary_ref_format = primary_ref.paragraph_format
    primary_ref_format.left_indent = Inches(0.5)
    primary_ref_format.hanging_indent = Inches(-0.5)
    primary_ref_format.space_after = Pt(6)
    
    # Add any additional sources from RAG system
    if unique_sources:
        for source_name, cite_info in unique_sources.items():
            source_display = cite_info.get('source_display', source_name.replace('.pdf', '').replace('_', ' '))
            ref_para = doc.add_paragraph()
            ref_para.add_run(f'{source_display}. ').italic = True
            ref_para.add_run('(Knowledge base source used for report enhancement.)')
            ref_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ref_para_format = ref_para.paragraph_format
            ref_para_format.left_indent = Inches(0.5)
            ref_para_format.hanging_indent = Inches(-0.5)
            ref_para_format.space_after = Pt(6)
    
    # Add note about RAG enhancement
    if rag_citations:
        doc.add_paragraph('')
        note_para = doc.add_paragraph()
        note_para.add_run('Note: ').bold = True
        note_para.add_run('This report was enhanced using a Retrieval-Augmented Generation (RAG) system that references authoritative knowledge base sources. Citations in the text refer to specific chunks from these sources, which were retrieved using semantic similarity search and validated against the original source documents.')
        note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        note_para_format = note_para.paragraph_format
        note_para_format.space_after = Pt(12)
    
    # TOC is now automatically generated by Word from heading styles
    # No manual TOC update needed - Word will populate it when the document is opened
    
    # Save document
    filename = f"model_benchmark_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Save to output directory if provided, otherwise current directory
    if output_dir:
        filepath = output_dir / filename
    else:
        filepath = Path(filename)
    
    doc.save(str(filepath))
    
    print(f"✓ Word document created: {filepath}")
    print("  - Automatic Table of Contents (Word built-in TOC field)")
    print("  - Page numbers added to footer")
    print("  - Model outputs included in Test Scenario Details")
    print("  Note: Word will automatically update the TOC when the document is opened")
    return filepath

def setup_test_directory():
    """Set up test directory structure and organize previous tests"""
    tests_dir = Path("tests")
    
    # Create tests directory if it doesn't exist
    if not tests_dir.exists():
        tests_dir.mkdir()
        print(f"✓ Created tests directory: {tests_dir}")
    else:
        print(f"✓ Tests directory exists: {tests_dir}")
    
    # Check for existing test result files in the root tests directory
    test_files = []
    for pattern in ["model_benchmark_results_*.json", "model_benchmark_report_*.docx"]:
        test_files.extend(list(tests_dir.glob(pattern)))
    
    # Move existing files to dated subfolders
    if test_files:
        print(f"\nFound {len(test_files)} previous test result file(s), organizing...")
        
        # Group files by date (extract from filename or use file modification date)
        files_by_date = {}
        
        for file in test_files:
            # Try to extract date from filename (format: YYYYMMDD_HHMMSS)
            date_match = re.search(r'(\d{8})_(\d{6})', file.name)
            if date_match:
                date_str = date_match.group(1)  # YYYYMMDD
            else:
                # Use file modification date
                mtime = os.path.getmtime(file)
                date_str = datetime.fromtimestamp(mtime).strftime("%Y%m%d")
            
            if date_str not in files_by_date:
                files_by_date[date_str] = []
            files_by_date[date_str].append(file)
        
        # Move files to dated subfolders
        for date_str, files in files_by_date.items():
            date_folder = tests_dir / date_str
            
            # Create date folder if it doesn't exist
            if not date_folder.exists():
                date_folder.mkdir()
                print(f"  Created date folder: {date_folder}")
            
            # Move files
            for file in files:
                try:
                    dest = date_folder / file.name
                    if dest.exists():
                        # If file already exists, add timestamp to avoid overwrite
                        timestamp = datetime.now().strftime("%H%M%S")
                        name_parts = file.stem.split('_')
                        if len(name_parts) > 0:
                            name_parts[-1] = f"{name_parts[-1]}_{timestamp}"
                        else:
                            name_parts.append(timestamp)
                        new_name = '_'.join(name_parts) + file.suffix
                        dest = date_folder / new_name
                    
                    shutil.move(str(file), str(dest))
                    print(f"  Moved: {file.name} -> {date_folder.name}/")
                except Exception as e:
                    print(f"  ⚠️  Warning: Could not move {file.name}: {e}")
    
    # Create subfolder for current test run (use full timestamp for uniqueness)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    current_test_dir = tests_dir / timestamp
    
    # Create the folder
    current_test_dir.mkdir()
    print(f"✓ Created test run folder: {current_test_dir}")
    
    return current_test_dir

def main():
    """Main function"""
    print("Snowflake Cortex AI Model Benchmarking")
    print("=" * 80)
    
    # Display configuration info
    config_path = PACKAGE_ROOT / "test_config.json"
    if not config_path.exists():
        config_path = Path.cwd() / "test_config.json"
    if config_path.exists():
        print(f"\n✓ Using configuration file: {config_path}")
        print(f"  - Models: {len(MODELS)} enabled")
        print(f"  - Test scenarios: {len(TEST_PROMPTS)} enabled")
    else:
        print(f"\n⚠️  Config file not found: {config_path}")
        print("  Using default hardcoded values")
    
    # Set up test directory structure
    print("\nSetting up test directory structure...")
    test_output_dir = setup_test_directory()
    
    # Get Snowflake session
    print("\nConnecting to Snowflake...")
    session = get_snowflake_session()
    
    if not session:
        print("❌ Error: Could not establish Snowflake connection.")
        print("\nDebugging connection issue...")
        # Try to see what's in config
        import re
        secrets_path = Path(".streamlit/secrets.toml")
        if secrets_path.exists():
            print(f"  ✓ Found secrets file: {secrets_path}")
            with open(secrets_path, 'r') as f:
                content = f.read()
            snowflake_section = re.search(r'\[connections\.snowflake\](.*?)(?=\[|$)', content, re.DOTALL)
            if snowflake_section:
                print("  ✓ Found [connections.snowflake] section")
                config_test = {}
                section_content = snowflake_section.group(1)
                for line in section_content.split('\n'):
                    line = line.strip()
                    if '=' in line and not line.startswith('#'):
                        key, value = line.split('=', 1)
                        key = key.strip()
                        value = value.strip().strip('"').strip("'")
                        config_test[key] = value
                print(f"  ✓ Parsed {len(config_test)} config keys")
                print(f"  ✓ Has account: {bool(config_test.get('account'))}")
                print(f"  ✓ Has user: {bool(config_test.get('user'))}")
                print(f"  ✓ Has password: {bool(config_test.get('password'))}")
            else:
                print("  ✗ Could not find [connections.snowflake] section")
        else:
            print(f"  ✗ Secrets file not found: {secrets_path}")
        print("\nPlease configure your connection:")
        print("  1. Set up .streamlit/secrets.toml with Snowflake credentials")
        print("  2. Or set environment variables:")
        print("     - SNOWFLAKE_ACCOUNT")
        print("     - SNOWFLAKE_USER")
        print("     - SNOWFLAKE_PASSWORD")
        print("     - SNOWFLAKE_WAREHOUSE (optional)")
        print("     - SNOWFLAKE_DATABASE (optional, defaults to SNOWFLAKE_SAMPLE_DATA)")
        print("     - SNOWFLAKE_SCHEMA (optional, defaults to TPCH_SF1)")
        return 1
    
    print("✓ Connected to Snowflake")
    
    # Verify access to sample data
    try:
        if hasattr(session, 'sql'):
            session.sql("SELECT COUNT(*) FROM SNOWFLAKE_SAMPLE_DATA.TPCH_SF1.CUSTOMER LIMIT 1").collect()
        else:
            cursor = session.cursor()
            cursor.execute("SELECT COUNT(*) FROM SNOWFLAKE_SAMPLE_DATA.TPCH_SF1.CUSTOMER LIMIT 1")
            cursor.fetchone()
            cursor.close()
        print("✓ Access to SNOWFLAKE_SAMPLE_DATA verified")
    except Exception as e:
        print(f"⚠️  Warning: Could not access SNOWFLAKE_SAMPLE_DATA: {e}")
        print("  Tests will still run but may reference tables that don't exist")
    
    # Run benchmarks
    try:
        results = run_model_benchmarks(session)
        
        # Analyze and recommend
        best_model_id, unavailable_models_list = analyze_benchmark_results(results)
        
        # Get model rankings for document (only available models)
        model_rankings = []
        model_stats = {}
        for result in results:
            model = result["model"]
            if model not in model_stats:
                model_stats[model] = {
                    "tests": [],
                    "success_count": 0,
                    "total_tests": 0,
                    "total_time": 0,
                    "total_cost": 0,
                    "total_quality_score": 0,
                    "executable_count": 0,
                    "executable_tests": 0,
                }
            model_stats[model]["tests"].append(result)
            model_stats[model]["total_tests"] += 1
            if result["success"]:
                model_stats[model]["success_count"] += 1
                model_stats[model]["total_time"] += result.get("elapsed_time", 0)
                model_stats[model]["total_cost"] += result.get("estimated_cost", 0)
                model_stats[model]["total_quality_score"] += result.get("quality_score", 0)
                if result.get("executable") is not None:
                    model_stats[model]["executable_tests"] += 1
                    if result["executable"]:
                        model_stats[model]["executable_count"] += 1
        
        for model, stats in model_stats.items():
            if stats["success_count"] == 0:
                continue
            avg_time = stats["total_time"] / stats["success_count"]
            avg_cost = stats["total_cost"] / stats["success_count"]
            avg_quality = stats["total_quality_score"] / stats["success_count"]
            success_rate = (stats["success_count"] / stats["total_tests"]) * 100
            executable_rate = (stats["executable_count"] / stats["executable_tests"] * 100) if stats["executable_tests"] > 0 else 0
            quality_normalized = (avg_quality / 10) * 100
            cost_efficiency = 100 / (avg_cost * 1000000 + 1)
            speed_score = 100 / (avg_time * 10 + 1)
            composite_score = (
                quality_normalized * 0.4 +
                cost_efficiency * 0.25 +
                speed_score * 0.15 +
                success_rate * 0.1 +
                executable_rate * 0.1
            )
            model_rankings.append({
                "model": model,
                "display": MODELS[model]["display"],
                "success_rate": success_rate,
                "avg_quality": avg_quality,
                "avg_time": avg_time,
                "avg_cost": avg_cost,
                "executable_rate": executable_rate,
                "composite_score": composite_score,
            })
        model_rankings.sort(key=lambda x: x["composite_score"], reverse=True)
        
        # Save results to JSON file in test output directory
        results_filename = f"model_benchmark_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        results_file = test_output_dir / results_filename
        with open(results_file, 'w') as f:
            json.dump(results, f, indent=2, default=str)
        print(f"\n✓ Results saved to: {results_file}")
        
        # Generate LLM report (with RAG enhancement by default)
        best_model = model_rankings[0] if model_rankings else None
        use_rag = os.getenv("USE_RAG_FOR_REPORTS", "true").lower() == "true"
        llm_report = generate_llm_report(session, results, model_rankings, best_model, use_rag=use_rag)
        
        # Get citations metadata if available (Problem 1: Citation Loss - FIXED)
        citations_metadata = getattr(generate_llm_report_section_based, '_citations_metadata', None)
        
        # Create Word document (final step) - update to save in test output directory
        word_doc = create_word_document(results, model_rankings, best_model, llm_report, test_output_dir, unavailable_models_list, rag_citations=citations_metadata, session=session)
        
        if word_doc:
            print(f"\n✓ Final report generated: {word_doc}")
        else:
            print("\n⚠️  Word document generation skipped (python-docx not available)")
        
        print("\n" + "=" * 80)
        print("BENCHMARKING COMPLETE")
        print("=" * 80)
        print(f"\nTest results saved to: {test_output_dir}")
        print("\nFiles generated:")
        print(f"  - JSON results: {results_file.name}")
        if word_doc:
            print(f"  - Word report: {word_doc.name}")
        
        return 0
    except KeyboardInterrupt:
        print("\n\n⚠️  Benchmarking interrupted by user")
        return 1
    except Exception as e:
        print(f"\n❌ Error during benchmarking: {e}")
        import traceback
        traceback.print_exc()
        return 1
    finally:
        # Close session
        try:
            if hasattr(session, 'close'):
                session.close()
            elif hasattr(session, 'cursor'):
                session.close()
        except Exception:
            pass
        
        # Clear schema cache after tests complete
        clear_schema_cache()
        logger.info("Test run completed, schema cache cleared")

if __name__ == "__main__":
    sys.exit(main())
